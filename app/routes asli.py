from flask import render_template, url_for, flash, redirect, request, send_from_directory, session, abort, send_file, abort, make_response, jsonify
from app import app, db, login_manager
from app.models import  User, Message, Siswa, Sekolah, Ekstrakurikuler, NilaiAkhir, Mapel, Kelas, RekapAbsensi, Guru, assign_wali_kelas, is_wali_kelas, AbsensiHarian, OrangTua, KegiatanIndustri, Semester,Pengajaran
from app.forms import  LoginForm, RegistrationForm, GuruForm, AddSekolahForm, EditGuruForm, EditSekolahForm, SiswaForm, KelasForm, MapelForm, NilaiAkhirForm, LaporanKelasForm, AbsensiHarianForm
from werkzeug.security import generate_password_hash
from flask_login import login_user, current_user, logout_user, login_required
from werkzeug.utils import secure_filename
from itertools import groupby
from operator import attrgetter
from zipfile import ZipFile
from docx import Document
from docx.shared import Pt, Cm  # Import Pt for point size
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # Import enum for paragraph alignment
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
from sqlalchemy import func
from functools import wraps
from collections import defaultdict
from sqlalchemy import create_engine, desc
from sqlalchemy.orm import sessionmaker
import pdfkit

from datetime import datetime
import io
from io import BytesIO
from urllib.parse import quote_plus
import tempfile
import pandas as pd
import os
import requests
import zipfile
import shutil


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if current_user.role != 'admin':
            abort(403)  # Forbidden
        return f(*args, **kwargs)
    return decorated_function

@app.errorhandler(403)
def forbidden(error):
    return render_template('403.html'), 403

@app.route('/dashboard')
def dashboard():
    kelas = Kelas.query.order_by(Kelas.nama_kelas.asc()).all()
    return render_template('dashboard.html', kelas=kelas)

@app.route('/bulk_delete_kelas')
def bulk_delete_kelas():
    kelas = Kelas.query.order_by(Kelas.nama_kelas.asc()).all() 
    return render_template('bulk_delete_kelas.html', kelas=kelas)   

@app.route('/bulk_delete_nilai_akhir', methods=['POST'])
def bulk_delete_nilai_akhir():
    kelas_id = request.form.get('kelas')
    if kelas_id:
        try:
            NilaiAkhir.query.filter_by(id_kelas=kelas_id).delete()
            db.session.commit()
            flash('Nilai akhir untuk kelas berhasil dihapus.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Terjadi kesalahan saat menghapus nilai akhir: {e}', 'danger')
    else:
        flash('Kelas tidak ditemukan.', 'danger')
    return redirect(url_for('bulk_delete_kelas'))


@app.route('/download_excel_siswa', methods=['POST'])
def download_excel_siswa():
    kelas_id = request.form['kelas']
    kelas = Kelas.query.get(kelas_id)
    
    wb = Workbook()
    ws = wb.active
    ws.append(['nama_siswa', 'nama_mapel', 'nama_guru', 'nilai', 'capaian_kompetensi','semester'])

    nilai_akhir_siswa = (
        db.session.query(NilaiAkhir, Siswa, Mapel, Guru)
        .join(Siswa)
        .join(Mapel)
        .join(Guru)
        .filter(NilaiAkhir.id_kelas == kelas_id)
        .order_by(Mapel.nama_mapel, Guru.nama_guru)
        .all()
    )

    for nilai_akhir, siswa, mapel, guru in nilai_akhir_siswa:
        ws.append([siswa.nama, mapel.nama_mapel, guru.nama_guru, '', '', ''])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"{kelas.nama_kelas}.xlsx")


@app.route('/download_excel_absensi', methods=['POST'])
def download_excel_absensi():
    kelas_id = request.form['kelas']
    kelas = Kelas.query.get(kelas_id)
    siswa = Siswa.query.filter_by(id_kelas=kelas_id).all()

    wb = Workbook()
    ws = wb.active
    ws.append(['nama_siswa', 'tahun_ajaran', 'total_sakit', 'total_izin', 'total_tanpa_keterangan', 'semester'])

    for s in siswa:
        ws.append([s.nama, '', '', '', '', ''])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"rekap_absensi_{kelas.nama_kelas}.xlsx")

@app.route('/download_excel_kegiatan_industri', methods=['POST'])
def download_excel_kegiatan_industri():
    kelas_id = request.form['kelas']
    kelas = Kelas.query.get(kelas_id)
    siswa = Siswa.query.filter_by(id_kelas=kelas_id).all()

    # Check if the selected class is class XI
    if not kelas.nama_kelas.lower().startswith('xi'):
        return "Only class XI students are allowed", 400

    wb = Workbook()
    ws = wb.active
    ws.append(['nama_siswa','tahun_ajaran','semester','mitra_induka', 'lokasi', 'keterangan' ])

    for s in siswa:
        # Assuming you have a relationship or method to get kegiatan_industri for a student
        kegiatan_industri = s.kegiatan_industri
        ws.append([s.nama, '', '', '', '', ''])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"rekap_kegiatan_industri_{kelas.nama_kelas}.xlsx")


@app.route('/download_siswa', methods=['POST'])
def download_siswa():
    kelas_id = request.form['kelas']
    kelas = Kelas.query.get(kelas_id)
    siswa_list = Siswa.query.filter_by(id_kelas=kelas_id).all()
    mapel_list = Mapel.query.filter_by(id_kelas=kelas_id).all()

    # Dictionary to group data by teacher and subject
    grouped_data = {}

    for siswa in siswa_list:
        for mapel in mapel_list:
            guru = Guru.query.get(mapel.id_guru)
            if not guru or not mapel:
                continue
            guru_mapel_key = (guru.nama_guru, mapel.nama_mapel)

            # Fetch the student's grade for the subject if it exists
            nilai_akhir = next((n for n in siswa.nilai_akhir if n.id_mapel == mapel.id_mapel), None)
            nilai = nilai_akhir.nilai if nilai_akhir else ''
            capaian_kompetensi = nilai_akhir.capaian_kompetensi if nilai_akhir else ''

            if guru_mapel_key not in grouped_data:
                grouped_data[guru_mapel_key] = []

            grouped_data[guru_mapel_key].append((siswa.nama, nilai, capaian_kompetensi))

    wb = Workbook()
    ws = wb.active
    ws.append(['nama_siswa','tahun_ajaran','semester', 'nama_mapel', 'nama_guru', 'nilai', 'capaian_kompetensi'])

    for (nama_guru, nama_mapel), entries in grouped_data.items():
        for entry in entries:
            nama_siswa, nilai, capaian_kompetensi = entry
            ws.append([nama_siswa,'','', nama_mapel, nama_guru, '', ''])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"{kelas.nama_kelas}.xlsx")


@app.route("/")
def index():
    return redirect(url_for('login'))


@app.route("/login", methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user)
            # Periksa role pengguna dan arahkan ke dashboard yang sesuai
            if user.role == 'admin':
                return redirect(url_for('admin_dashboard'))
            elif user.role == 'guru':
                return redirect(url_for('guru_dashboard'))
            else:
                # Jika role tidak dikenal, arahkan ke halaman utama atau halaman error
                return redirect(url_for('home'))
        else:
            flash('Login Gagal. Silakan periksa username dan password', 'danger')
    return render_template('login.html', title='Login', form=form)


@app.route("/logout")
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route("/register", methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('home'))  # Pengguna yang sudah login tidak perlu registrasi

    form = RegistrationForm()
    if form.validate_on_submit():
        hashed_password = generate_password_hash(form.password.data)
        user = User(username=form.username.data, password=hashed_password, role=form.role.data)
        db.session.add(user)
        db.session.commit()
        flash('Your account has been created! You are now able to log in', 'success')
        return redirect(url_for('login'))

    return render_template('register.html', title='Register', form=form)

@app.route('/create', methods=['POST'])
def create():
    new_message = request.form.get('new_message')
    if new_message:
        message = Message(message=new_message)
        db.session.add(message)
        db.session.commit()
    return redirect(url_for('dashboard'))

@app.route('/delete/<int:id>', methods=['GET', 'POST'])
def delete(id):
    user = User.query.get(id)
    
    db.session.delete(user)
    db.session.commit()
    flash('User berhasil dihapus', 'danger')
    return redirect('dashboard')

@app.route("/admin/dashboard")
@login_required
def admin_dashboard():
    if current_user.role != 'admin':
        return redirect(url_for('home'))  # Arahkan non-admin ke halaman utama
    # Tampilkan konten dashboard admin
    return render_template('dashboard.html')

@app.route("/guru/dashboard")
@login_required
def guru_dashboard():
    if current_user.role != 'guru':
        return redirect(url_for('home'))  # Arahkan non-guru ke halaman utama
    # Tampilkan konten dashboard guru
    return render_template('dashboard.html')

@app.route("/add-guru", methods=['GET', 'POST'])
@login_required
@admin_required
def add_guru():
    form = GuruForm()
    if form.validate_on_submit():
        guru = Guru(nama_guru=form.nama_guru.data, user_id=form.user_id.data)
        db.session.add(guru)
        db.session.commit()
        flash('Guru baru telah ditambahkan', 'success')
        return redirect(url_for('list_guru'))  # Sesuaikan dengan fungsi tujuan

    return render_template('add_guru.html', title='Tambah Guru', form=form)

@app.route('/edit-guru/<int:id>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_guru(id):
    guru = Guru.query.get_or_404(id)
    form = EditGuruForm(obj=guru)
    if form.validate_on_submit():
        guru.nama_guru = form.nama_guru.data
        db.session.commit()
        flash('Data guru telah diperbarui', 'success')
        return redirect(url_for('list_guru'))
    return render_template('edit_guru.html', form=form)

@app.route('/delete-guru/<int:id>', methods=['POST'])
@login_required
@admin_required
def delete_guru(id):
    guru = Guru.query.get_or_404(id)
    db.session.delete(guru)
    db.session.commit()
    flash('Guru telah dihapus', 'success')
    return redirect(url_for('list_guru'))

@app.route('/list-guru')
@login_required

def list_guru():
    guru = Guru.query.all()
    return render_template('list_guru.html', guru=guru)


@app.route('/add-sekolah', methods=['GET', 'POST'])
@login_required
@admin_required
def add_sekolah():
    form = AddSekolahForm()
    if form.validate_on_submit():
        sekolah = Sekolah(nama_sekolah=form.nama_sekolah.data, alamat_sekolah=form.alamat_sekolah.data)
        db.session.add(sekolah)
        db.session.commit()
        flash('Sekolah baru telah ditambahkan', 'success')
        return redirect(url_for('list_sekolah'))  # Sesuaikan dengan fungsi tujuan

    return render_template('add_sekolah.html', title='Tambah Sekolah', form=form)

@app.route('/edit-sekolah/<int:id>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_sekolah(id):
    sekolah = Sekolah.query.get_or_404(id)
    form = EditSekolahForm(obj=sekolah)
    if form.validate_on_submit():
        sekolah.nama_sekolah = form.nama_sekolah.data
        sekolah.alamat_sekolah = form.alamat_sekolah.data
        db.session.commit()
        flash('Data sekolah telah diperbarui', 'success')
        return redirect(url_for('list_sekolah'))
    return render_template('edit_sekolah.html', form=form)

@app.route('/delete-sekolah/<int:id>', methods=['POST'])
@login_required
@admin_required
def delete_sekolah(id):
    sekolah = Sekolah.query.get_or_404(id)
    db.session.delete(sekolah)
    db.session.commit()
    flash('Sekolah telah dihapus', 'success')
    return redirect(url_for('list_sekolah'))


@app.route('/list-sekolah')
@login_required
@admin_required
def list_sekolah():
    sekolah = Sekolah.query.all()
    return render_template('list_sekolah.html', sekolah=sekolah)

@app.route('/add-siswa', methods=['GET', 'POST'])
@login_required
@admin_required
def add_siswa():
    form = SiswaForm()
    if form.validate_on_submit():
        siswa = Siswa(nisn=form.nisn.data, nama=form.nama.data, alamat=form.alamat.data, id_kelas=form.id_kelas.data)
        db.session.add(siswa)
        db.session.commit()
        flash('Siswa baru telah ditambahkan', 'success')
        return redirect(url_for('list_siswa'))
    return render_template('add_siswa.html', form=form)

@app.route('/edit-siswa/<int:id_siswa>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_siswa(id_siswa):
    siswa = Siswa.query.get_or_404(id_siswa)
    form = SiswaForm(obj=siswa)
    if form.validate_on_submit():
        siswa.nisn = form.nisn.data
        siswa.nama = form.nama.data
        siswa.alamat = form.alamat.data
        siswa.id_kelas = form.id_kelas.data
        db.session.commit()
        flash('Data siswa telah diperbarui', 'success')
        return redirect(url_for('list_siswa'))
    return render_template('edit_siswa.html', form=form)

@app.route('/delete-siswa/<int:id_siswa>', methods=['POST'])
@login_required
@admin_required
def delete_siswa(id_siswa):
    siswa = Siswa.query.get_or_404(id_siswa)
    db.session.delete(siswa)
    db.session.commit()
    flash('Siswa telah dihapus', 'success')
    return redirect(url_for('list_siswa'))

@app.route('/list-siswa')
@login_required

def list_siswa():
    kelas = request.args.get('kelas')
    semua_kelas = Kelas.query.with_entities(Kelas.nama_kelas).distinct().all()
    if kelas:
        siswa = Siswa.query.join(Kelas).filter(Kelas.nama_kelas == kelas).all()
    else:
        siswa = Siswa.query.all()
    return render_template('list_siswa.html', siswa=siswa, semua_kelas=semua_kelas)

@app.route('/download-siswa-excel')
@login_required
def download_siswa_excel():
    kelas = request.args.get('kelas')
    if kelas:
        siswa = Siswa.query.join(Kelas).filter(Kelas.nama_kelas == kelas).all()
        file_name = f"data_siswa_{kelas}.xlsx"
    else:
        siswa = Siswa.query.all()
        file_name = "data_siswa.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Data Siswa"

    headers = ['NISN', 'Nama', 'Alamat', 'Nama Kelas', 'Tingkat']
    for col_num, header in enumerate(headers, 1):
        col_letter = get_column_letter(col_num)
        ws['{}1'.format(col_letter)] = header
        ws.column_dimensions[col_letter].width = 15

    for row_num, data in enumerate(siswa, 2):
        ws.cell(row=row_num, column=1, value=data.nisn)
        ws.cell(row=row_num, column=2, value=data.nama)
        ws.cell(row=row_num, column=3, value=data.alamat)
        ws.cell(row=row_num, column=4, value=data.kelas.nama_kelas)
        ws.cell(row=row_num, column=5, value=data.kelas.tingkat)

    excel_data = io.BytesIO()
    wb.save(excel_data)
    excel_data.seek(0)

    return send_file(
        excel_data, 
        as_attachment=True,
        download_name=file_name,  # Correct argument for specifying the download file name
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@app.route('/add-kelas', methods=['GET', 'POST'])
@login_required
@admin_required
def add_kelas():
    form = KelasForm()
    # Isi pilihan id_guru dengan daftar guru
    form.id_guru.choices = [(guru.id_guru, guru.nama_guru) for guru in Guru.query.all()]
    
    if form.validate_on_submit():
        kelas = Kelas(nama_kelas=form.nama_kelas.data, tingkat=form.tingkat.data, id_guru=form.id_guru.data)
        db.session.add(kelas)
        db.session.commit()
        flash('Kelas baru telah ditambahkan', 'success')
        return redirect(url_for('list_kelas'))

    return render_template('add_kelas.html', form=form)

@app.route('/edit-kelas/<int:id_kelas>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_kelas(id_kelas):
    kelas = Kelas.query.get_or_404(id_kelas)
    form = KelasForm(obj=kelas)
    form.id_guru.choices = [(guru.id_guru, guru.nama_guru) for guru in Guru.query.all()]
    if form.validate_on_submit():
        form.populate_obj(kelas)
        db.session.commit()
        flash('Kelas berhasil diubah', 'success')
        return redirect(url_for('list_kelas'))
    return render_template('edit_kelas.html', form=form, kelas=kelas)

@app.route('/delete-kelas/<int:id_kelas>', methods=['POST'])
@login_required
@admin_required
def delete_kelas(id_kelas):
    kelas = Kelas.query.get_or_404(id_kelas)
    db.session.delete(kelas)
    db.session.commit()
    flash('Kelas telah dihapus', 'success')
    return redirect(url_for('list_kelas'))

@app.route('/list-kelas')
@login_required
def list_kelas():
    
    kelas = Kelas.query.all()
    return render_template('list_kelas.html', kelas=kelas)


@app.route('/add-mapel', methods=['GET', 'POST'])
@login_required
@admin_required
def add_mapel():
    form = MapelForm()
    form.id_kelas.choices = [(kelas.id_kelas, kelas.nama_kelas) for kelas in Kelas.query.all()]
    form.id_guru.choices = [(guru.id_guru, guru.nama_guru) for guru in Guru.query.all()]  # Menambahkan pilihan guru ke dalam form
    if form.validate_on_submit():
        mapel = Mapel(
            nama_mapel=form.nama_mapel.data,
            deskripsi=form.deskripsi.data,
            jumlah_jam=form.jumlah_jam.data,
            kategori=form.kategori.data,
            id_kelas=form.id_kelas.data,
            id_guru=form.id_guru.data  # Menyimpan id guru yang dipilih dari form
        )
        db.session.add(mapel)
        db.session.commit()
        flash('Mata pelajaran berhasil ditambahkan!')
        return redirect(url_for('list_mapel'))
    return render_template('add_mapel.html', form=form)

@app.route('/edit-mapel/<int:id_mapel>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_mapel(id_mapel):
    mapel = Mapel.query.get_or_404(id_mapel)
    form = MapelForm(obj=mapel)
    form.id_kelas.choices = [(kelas.id_kelas, kelas.nama_kelas) for kelas in Kelas.query.all()]
    if form.validate_on_submit():
        mapel.nama_mapel = form.nama_mapel.data
        mapel.deskripsi = form.deskripsi.data
        mapel.jumlah_jam = form.jumlah_jam.data
        mapel.kategori = form.kategori.data
        mapel.id_kelas = form.id_kelas.data
        db.session.commit()
        flash('Mata pelajaran berhasil diperbarui!')
        return redirect(url_for('list_mapel'))
    return render_template('edit_mapel.html', form=form, mapel=mapel)

@app.route('/delete-mapel/<int:id_mapel>', methods=['POST'])
@login_required
@admin_required
def delete_mapel(id_mapel):
    mapel = Mapel.query.get_or_404(id_mapel)
    db.session.delete(mapel)
    db.session.commit()
    flash('Mata pelajaran berhasil dihapus!')
    return redirect(url_for('list_mapel'))

@app.route('/list-mapel')
@login_required
def list_mapel():
    mapel_list = Mapel.query.all()
    return render_template('list_mapel.html', mapel_list=mapel_list)


def read_config():
    """Baca tanggal dari file konfigurasi."""
    config = {}
    try:
        with open('config.txt', 'r') as f:
            for line in f:
                key, value = line.strip().split('=')
                config[key] = datetime.strptime(value, '%Y-%m-%d')
    except Exception as e:
        print(f"Error reading config file: {e}")
    return config

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        # Ambil tanggal dari formulir
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')

        # Simpan tanggal ke file (atau basis data)
        with open('config.txt', 'w') as f:
            f.write(f'start_date={start_date}\nend_date={end_date}')

        return redirect(url_for('upload'))

    # Tampilkan halaman pengaturan
    return render_template('settings.html')



@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    config = read_config()
    start_date = config.get('start_date')
    end_date = config.get('end_date')

    # Periksa tanggal
    if not (start_date <= datetime.now() <= end_date):
        return 'Akses hanya diizinkan antara {} dan {}'.format(start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d'))

    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part in the request', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        if file:
            df = pd.read_excel(file, engine='openpyxl')

            # Rename columns if necessary
            df.rename(columns={
                'Nama Siswa': 'nama_siswa',
                'Tahun Ajaran': 'tahun_ajaran',
                'Semester': 'semester',
                'Nama Mapel': 'nama_mapel',
                'Nama Guru': 'nama_guru',
                'Nilai': 'nilai',
                'Capaian Kompetensi': 'capaian_kompetensi'
            }, inplace=True)

            # Print columns for debugging
            print(df.columns)
            flash(f"Columns in the uploaded file: {', '.join(df.columns)}", 'info')

            # Check if required columns exist
            required_columns = ['nama_siswa', 'tahun_ajaran', 'semester', 'nama_mapel', 'nama_guru', 'nilai', 'capaian_kompetensi']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                flash(f"Missing columns in the uploaded file: {', '.join(missing_columns)}", 'error')
                return redirect(request.url)
            
            session['df'] = df.to_html(classes='table table-striped', index=False, border=0)  # Save DataFrame as HTML

    return render_template('add_nilaiakhir.html', data=session.get('df'))  # Render DataFrame as HTML table


@app.route('/form_page', methods=['GET'])
@login_required
def form_page():
    # Hapus data yang diunggah sebelumnya dari session
    session.pop('df', None)

    # Render template untuk form upload
    return render_template('add_nilaiakhir.html')

"""
@app.route('/submit', methods=['POST'])
@login_required
def submit():
    df = pd.read_html(session['df'])[0]
    error_messages = []

    for index, row in df.iterrows():
        siswa = Siswa.query.filter_by(nama=row['nama_siswa']).first()
        mapel = Mapel.query.filter_by(nama_mapel=row['nama_mapel']).join(Kelas).filter(Kelas.id_kelas == siswa.id_kelas).first() if siswa else None
        guru = Guru.query.filter_by(nama_guru=row['nama_guru']).first()
        semester = Semester.query.filter_by(semester=row['semester']).first()

        print(f"Checking for Siswa: {siswa}, Mapel: {mapel}, Guru: {guru}, Semester: {semester}")

        if not siswa:
            error_messages.append(f"Siswa '{row['nama_siswa']}' tidak ditemukan.")
            continue
        if not mapel:
            error_messages.append(f"Mata Pelajaran '{row['nama_mapel']}' untuk kelas siswa '{row['nama_siswa']}' tidak ditemukan.")
            continue
        if not guru:
            error_messages.append(f"Guru '{row['nama_guru']}' tidak ditemukan.")
            continue
        if not semester:
            error_messages.append(f"Semester '{row['semester']}' tidak ditemukan.")
            continue

        if siswa.id_kelas != mapel.id_kelas:
            error_messages.append(f"Siswa '{siswa.nama}' tidak terdaftar di kelas untuk mata pelajaran '{mapel.nama_mapel}'.")
            continue

        existing_entry = NilaiAkhir.query.filter_by(
            id_siswa=siswa.id_siswa, 
            id_mapel=mapel.id_mapel, 
            id_guru=guru.id_guru,
            id_kelas=siswa.id_kelas,
            id_semester=semester.id
        ).first()

        if existing_entry:
            error_messages.append(f"Data untuk siswa '{siswa.nama}', mata pelajaran '{mapel.nama_mapel}', guru '{guru.nama_guru}', kelas '{siswa.kelas.nama_kelas}', dan semester '{semester.semester}' sudah ada.")
            continue

        nilai_akhir = NilaiAkhir(
            id_siswa=siswa.id_siswa,
            id_mapel=mapel.id_mapel,
            nilai=row['nilai'],
            capaian_kompetensi=row['capaian_kompetensi'],
            id_guru=guru.id_guru,
            id_kelas=siswa.id_kelas,
            id_semester=semester.id
        )
        db.session.add(nilai_akhir)

    db.session.commit()

    if error_messages:
        for msg in error_messages:
            flash(msg, 'error')
    else:
        flash('Data submitted successfully', 'success')

    if not error_messages:
        flash('Data submitted successfully', 'success')
        session.pop('df', None)

    return redirect(url_for('form_page'))
"""

@app.route('/submit', methods=['POST'])
@login_required
def submit():
    df = pd.read_html(session['df'])[0]
    error_messages = []

    for index, row in df.iterrows():
        siswa = Siswa.query.filter_by(nama=row['nama_siswa']).first()
        mapel = Mapel.query.filter_by(nama_mapel=row['nama_mapel']).join(Kelas).filter(Kelas.id_kelas == siswa.id_kelas).first() if siswa else None
        guru = Guru.query.filter_by(nama_guru=row['nama_guru']).first()
        semester = Semester.query.filter_by(tahun_ajaran=row['tahun_ajaran'], semester=row['semester']).first()

        if not siswa:
            error_messages.append(f"Siswa '{row['nama_siswa']}' tidak ditemukan.")
            continue
        if not mapel:
            error_messages.append(f"Mata Pelajaran '{row['nama_mapel']}' untuk kelas siswa '{row['nama_siswa']}' tidak ditemukan.")
            continue
        if not guru:
            error_messages.append(f"Guru '{row['nama_guru']}' tidak ditemukan.")
            continue
        if not semester:
            error_messages.append(f"Semester '{row['semester']}' tidak ditemukan.")
            continue

        if siswa.id_kelas != mapel.id_kelas:
            error_messages.append(f"Siswa '{siswa.nama}' tidak terdaftar di kelas untuk mata pelajaran '{mapel.nama_mapel}'.")
            continue

        existing_entry = NilaiAkhir.query.filter_by(
            id_siswa=siswa.id_siswa, 
            id_mapel=mapel.id_mapel, 
            id_guru=guru.id_guru,
            id_kelas=siswa.id_kelas,
            id_semester=semester.id
        ).first()

        if existing_entry:
            error_messages.append(f"Data untuk siswa '{siswa.nama}', mata pelajaran '{mapel.nama_mapel}', guru '{guru.nama_guru}', kelas '{siswa.kelas.nama_kelas}', dan semester '{semester.semester}' sudah ada.")
            continue

        nilai_akhir = NilaiAkhir(
            id_siswa=siswa.id_siswa,
            id_mapel=mapel.id_mapel,
            nilai=row['nilai'],
            capaian_kompetensi=row['capaian_kompetensi'],
            id_guru=guru.id_guru,
            id_kelas=siswa.id_kelas,
            id_semester=semester.id
        )
        db.session.add(nilai_akhir)

    db.session.commit()

    if error_messages:
        for msg in error_messages:
            flash(msg, 'error')
    else:
        flash('Data submitted successfully', 'success')

    if not error_messages:
        flash('Data submitted successfully', 'success')
        session.pop('df', None)

    return redirect(url_for('form_page'))


@app.route('/add-nilai-akhir', methods=['GET', 'POST'])
@login_required
def add_nilai_akhir():
    if request.method == 'POST':
        id_siswa = request.form.get('id_siswa')
        id_mapel = request.form.get('id_mapel')
        id_guru = request.form.get('id_guru')
        nilai = request.form.get('nilai')
        capaian_kompetensi = request.form.get('capaian_kompetensi')

        nilai_akhir = NilaiAkhir(id_siswa=id_siswa, id_mapel=id_mapel, id_guru=id_guru, nilai=nilai, capaian_kompetensi=capaian_kompetensi)
        db.session.add(nilai_akhir)
        db.session.commit()
        flash('Nilai akhir berhasil ditambahkan', 'success')
        return redirect(url_for('list_nilai_akhir'))

    return render_template('add_nilai_akhir.html')

@app.route('/edit-nilai-akhir/<int:id_nilai>', methods=['GET', 'POST'])
@login_required
def edit_nilai_akhir(id_nilai):
    nilai_akhir = NilaiAkhir.query.get_or_404(id_nilai)

    if request.method == 'POST':
        nilai_akhir.id_siswa = request.form.get('id_siswa')
        nilai_akhir.id_mapel = request.form.get('id_mapel')
        nilai_akhir.id_guru = request.form.get('id_guru')
        nilai_akhir.nilai = request.form.get('nilai')
        nilai_akhir.capaian_kompetensi = request.form.get('capaian_kompetensi')

        db.session.commit()
        flash('Nilai akhir berhasil diperbarui', 'success')
        return redirect(url_for('list_nilai_akhir'))

    return render_template('edit_nilai_akhir.html', nilai_akhir=nilai_akhir)

@app.route('/delete-nilai-akhir/<int:id_nilai>', methods=['POST'])
@login_required
@admin_required
def delete_nilai_akhir(id_nilai):
    nilai_akhir = NilaiAkhir.query.get_or_404(id_nilai)
    db.session.delete(nilai_akhir)
    db.session.commit()
    flash('Nilai akhir berhasil dihapus', 'success')
    return redirect(url_for('list_nilai_akhir'))


@app.route('/list-nilai-akhir')
def list_nilai_akhir():
    nilai_akhir_list = db.session.query(NilaiAkhir, Siswa, Mapel, Guru).join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa).join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel).join(Guru, NilaiAkhir.id_guru == Guru.id_guru).all()
    return render_template('list_nilai_akhir.html', nilai_akhir_list=nilai_akhir_list)

@app.route('/upload-rekap-absensi', methods=['GET', 'POST'])
@login_required
def upload_rekap_absensi():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part in the request', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        if file:
            df = pd.read_excel(file, engine='openpyxl')
            error_messages = []  # Untuk menyimpan pesan kesalahan
            for _, row in df.iterrows():
                siswa = Siswa.query.filter_by(nama=row['nama_siswa']).first()
                if not siswa:
                    # Rekam pesan kesalahan jika siswa tidak ditemukan
                    error_messages.append(f"Siswa dengan nama '{row['nama_siswa']}' tidak ditemukan.")
                    continue  # Lanjutkan ke baris berikutnya

                semester = Semester.query.filter_by(tahun_ajaran=row['tahun_ajaran'], semester=row['semester']).first()
                if not semester:
                    # Rekam pesan kesalahan jika semester tidak ditemukan
                    error_messages.append(f"Semester '{row['semester']}' untuk tahun ajaran '{row['tahun_ajaran']}' tidak ditemukan.")
                    continue  # Lanjutkan ke baris berikutnya

                rekap_absensi = RekapAbsensi(
                    id_siswa=siswa.id_siswa,
                    tahun_ajaran=row['tahun_ajaran'],
                    total_sakit=row['total_sakit'],
                    total_izin=row['total_izin'],
                    total_tanpa_keterangan=row['total_tanpa_keterangan'],
                    id_semester=semester.id
                )
                db.session.add(rekap_absensi)
            db.session.commit()
            flash('Data berhasil di-submit', 'success')
            # Tampilkan pesan kesalahan (jika ada)
            if error_messages:
                flash('Beberapa entri tidak bisa diproses:\n' + '\n'.join(error_messages), 'error')

            session['df'] = df.to_html(classes='table table-striped', index=False, border=0)

    return render_template('upload_rekap_absensi.html', data=session.get('df'))

@app.route('/upload-ekstrakurikuler', methods=['GET', 'POST'])
@login_required
def upload_ekstrakurikuler():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part in the request', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        if file:
            df = pd.read_excel(file, engine='openpyxl')

            # Check if required columns exist
            required_columns = ['nama_siswa', 'tahun_ajaran', 'semester', 'kegiatan', 'predikat', 'keterangan']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                flash(f"Missing columns in the uploaded file: {', '.join(missing_columns)}", 'error')
                return redirect(request.url)
            
            error_messages = []  # Untuk menyimpan pesan kesalahan
            for _, row in df.iterrows():
                siswa = Siswa.query.filter_by(nama=row['nama_siswa']).first()
                if not siswa:
                    error_messages.append(f"Siswa dengan nama '{row['nama_siswa']}' tidak ditemukan.")
                    continue  # Lanjutkan ke iterasi berikutnya

                semester = Semester.query.filter_by(tahun_ajaran=row['tahun_ajaran'], semester=row['semester']).first()
                if not semester:
                    error_messages.append(f"Semester '{row['semester']}' untuk tahun ajaran '{row['tahun_ajaran']}' tidak ditemukan.")
                    continue  # Lanjutkan ke iterasi berikutnya

                ekstrakurikuler = Ekstrakurikuler(
                    id_siswa=siswa.id_siswa,
                    kegiatan=row['kegiatan'],
                    predikat=row['predikat'],
                    keterangan=row['keterangan'],
                    id_semester=semester.id
                )
                db.session.add(ekstrakurikuler)
            db.session.commit()

            # Menampilkan pesan kesalahan atau sukses
            if error_messages:
                flash('Beberapa entri tidak bisa diproses:\n' + '\n'.join(error_messages), 'error')
            else:
                flash('Data ekstrakurikuler berhasil di-upload', 'success')

    return render_template('upload_ekstrakurikuler.html')

@app.route('/upload-kegiatan-industri', methods=['GET', 'POST'])
@login_required
def upload_kegiatan_industri():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part in the request', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        if file:
            df = pd.read_excel(file, engine='openpyxl')

            # Check if required columns exist
            required_columns = ['nama_siswa', 'tahun_ajaran', 'semester', 'mitra_induka', 'lokasi', 'keterangan']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                flash(f"Missing columns in the uploaded file: {', '.join(missing_columns)}", 'error')
                return redirect(request.url)
            
            error_messages = []  # Untuk menyimpan pesan kesalahan
            for _, row in df.iterrows():
                siswa = Siswa.query.filter_by(nama=row['nama_siswa']).first()
                if not siswa:
                    error_messages.append(f"Siswa dengan nama '{row['nama_siswa']}' tidak ditemukan.")
                    continue  # Lanjutkan ke baris berikutnya

                semester = Semester.query.filter_by(tahun_ajaran=row['tahun_ajaran'], semester=row['semester']).first()
                if not semester:
                    error_messages.append(f"Semester '{row['semester']}' untuk tahun ajaran '{row['tahun_ajaran']}' tidak ditemukan.")
                    continue  # Lanjutkan ke baris berikutnya

                kegiatan_industri = KegiatanIndustri(
                    id_siswa=siswa.id_siswa,
                    mitra_induka=row['mitra_induka'],
                    lokasi=row['lokasi'],
                    keterangan=row['keterangan'],
                    id_semester=semester.id
                )
                db.session.add(kegiatan_industri)
            db.session.commit()
            flash('Data berhasil di-submit', 'success')
            # Tampilkan pesan kesalahan (jika ada)
            if error_messages:
                flash('Beberapa entri tidak bisa diproses:\n' + '\n'.join(error_messages), 'error')

            session['df_kegiatan_industri'] = df.to_html(classes='table table-striped', index=False, border=0)

    return render_template('upload_kegiatan_industri.html', data=session.get('df_kegiatan_industri'))


# end upload modul

@app.route('/get-guru-by-kelas/<int:kelas_id>')
@login_required
def get_guru_by_kelas(kelas_id):
    # Query untuk mengambil guru yang mengajar di kelas yang dipilih
    guru = db.session.query(Guru).join(NilaiAkhir).join(Siswa).filter(Siswa.id_kelas == kelas_id).distinct().all()
    guru_list = [{'id': g.id_guru, 'name': g.nama_guru} for g in guru]
    return jsonify(guru_list)

"""
@app.route('/laporan-kelas', methods=['GET', 'POST'])
@login_required
def nilai_akhir():
    kelas = Kelas.query.order_by(Kelas.tingkat).all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran.asc(), Semester.semester.asc()).all()  # Menambahkan query untuk mengurutkan semester

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        semester_id = request.form.get('semester')  # Menangkap nilai semester yang dipilih

        hasil = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id)\
            .all()

        return render_template('laporan.html', hasil=hasil, kelas=kelas, semesters=semesters)

    else:
        return render_template('laporan.html', kelas=kelas, semesters=semesters)
"""
"""
@app.route('/laporan-kelas', methods=['GET', 'POST'])
@login_required
def nilai_akhir():
    kelas = Kelas.query.order_by(Kelas.tingkat).all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran, Semester.semester).all()

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        semester_id = request.form.get('semester')

        hasil = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id)\
            .all()

        nilai_bawah_76 = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id, NilaiAkhir.nilai < 76)\
            .all()

        return render_template('laporan.html', hasil=hasil, kelas=kelas, semesters=semesters, nilai_bawah_76=nilai_bawah_76)

    else:
        return render_template('laporan.html', kelas=kelas, semesters=semesters)

"""
"""
@app.route('/laporan-kelas', methods=['GET', 'POST'])
@login_required
def nilai_akhir():
    kelas = Kelas.query.order_by(Kelas.tingkat).all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran, Semester.semester).all()

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        semester_id = request.form.get('semester')

        hasil = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id)\
            .all()

        nilai_bawah_76 = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id, NilaiAkhir.nilai < 76)\
            .all()

        return render_template('laporan.html', hasil=hasil, kelas=kelas, semesters=semesters, nilai_bawah_76=nilai_bawah_76)

    else:
        return render_template('laporan.html', kelas=kelas, semesters=semesters)

@app.route('/update-nilai', methods=['POST'])
@login_required
def update_nilai():
    nilai_ids = request.form.getlist('nilai_id')
    nilai_values = request.form.getlist('nilai')

    for nilai_id, nilai in zip(nilai_ids, nilai_values):
        nilai_akhir = NilaiAkhir.query.get(nilai_id)
        nilai_akhir.nilai = nilai
        db.session.commit()

    return redirect(url_for('nilai_akhir'))
"""

@app.route('/laporan-kelas', methods=['GET', 'POST'])
@login_required
def nilai_akhir():
    kelas = Kelas.query.order_by(Kelas.tingkat).all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran, Semester.semester).all()

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        semester_id = request.form.get('semester')

        hasil = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id)\
            .all()

        nilai_bawah_76 = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
            .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
            .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
            .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
            .filter(Siswa.id_kelas == kelas_id, Guru.id_guru == guru_id, NilaiAkhir.id_semester == semester_id, NilaiAkhir.nilai < 66)\
            .all()

        return render_template('laporan.html', hasil=hasil, kelas=kelas, semesters=semesters, nilai_bawah_76=nilai_bawah_76)

    else:
        return render_template('laporan.html', kelas=kelas, semesters=semesters)

@app.route('/update-nilai', methods=['POST'])
@login_required
def update_nilai():
    token = request.form.get('token')
    if token != "123456":
        flash("Token tidak valid.", "danger")
        return redirect(url_for('nilai_akhir'))

    nilai_ids = request.form.getlist('nilai_id')
    nilai_values = request.form.getlist('nilai')

    for nilai_id, nilai in zip(nilai_ids, nilai_values):
        nilai_akhir = NilaiAkhir.query.get(nilai_id)
        nilai_akhir.nilai = nilai
        db.session.commit()

    flash("Nilai berhasil diperbarui.", "success")
    return redirect(url_for('nilai_akhir'))



@app.route('/edit-nilai/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_nilai(id):
    nilai_akhir = NilaiAkhir.query.get_or_404(id)
    if request.method == 'POST':
        nilai_akhir.nilai = request.form.get('nilai')
        db.session.commit()
        return redirect(url_for('nilai_akhir'))
    return render_template('edit_nilai.html', nilai_akhir=nilai_akhir)




@app.route('/laporan-kelas1', methods=['GET', 'POST'])
@login_required
def laporan():
    form = LaporanKelasForm()
    form.kelas.choices = [(kelas.id_kelas, kelas.nama_kelas) for kelas in Kelas.query.all()]
    if form.validate_on_submit():
        kelas = Kelas.query.get(form.kelas.data)
        mapel = Mapel.query.filter_by(id_kelas=kelas.id_kelas).all()
        guru = Guru.query.filter_by(id_guru=kelas.id_guru).first()
        return render_template('laporan_kelas.html', kelas=kelas, mapel=mapel, guru=guru)
    else:
        return render_template('laporan_form_kelas.html', form=form)



@app.route('/laporan-guru', methods=['GET'])
@login_required
def laporan_guru():
    # Fetch students who have chosen a religion
    students_with_religion = db.session.query(Siswa.nama).filter(Siswa.religion != None).all()
    students_with_religion = [student.nama for student in students_with_religion]

    # Ambil data dari database dan urutkan berdasarkan nama guru dan nama kelas
    data = db.session.query(Guru, Mapel, Kelas, Siswa, NilaiAkhir)\
        .join(Mapel, Guru.id_guru == Mapel.id_guru)\
        .join(Kelas, Mapel.id_kelas == Kelas.id_kelas)\
        .join(Siswa, Kelas.id_kelas == Siswa.id_kelas)\
        .outerjoin(NilaiAkhir, (NilaiAkhir.id_guru == Guru.id_guru) & (NilaiAkhir.id_mapel == Mapel.id_mapel) & (NilaiAkhir.id_siswa == Siswa.id_siswa))\
        .filter(NilaiAkhir.id_nilai == None)\
        .order_by(Guru.nama_guru, Kelas.nama_kelas)\
        .all()

    # Filter siswa agar hanya muncul di bawah guru agama yang mereka ikuti
    grouped_data = []
    for key, group in groupby(data, key=lambda x: (x.Guru.nama_guru, x.Kelas.nama_kelas)):
        group_list = list(group)
        filtered_group_list = []

        for item in group_list:
            if item.Mapel.nama_mapel in ['Pendidikan Agama Kristen dan Budi Pekerti',
                                        'Pendidikan Agama Islam dan Budi Pekerti',
                                        'Pendidikan Agama Katolik dan Budi Pekerti']:
                # Filter out students who have not chosen a religion
                if item.Siswa.nama in students_with_religion:
                    filtered_group_list.append(item)
            else:
                filtered_group_list.append(item)

        if filtered_group_list:
            grouped_data.append({
                'guru': key[0],
                'kelas': key[1],
                'mapel': filtered_group_list
            })

    # Render template dengan data yang sudah difilter
    return render_template('laporan_guru.html', grouped_data=grouped_data)



# Fungsi untuk mengatur gaya sel tabel


def atur_gaya_sel(sel, nama_font='Times New Roman', ukuran_font=12, tengah_align=False, justify_align=False):
    """
    Mengatur gaya untuk sel dalam dokumen.

    Parameters:
    - sel: Sel yang akan diatur gayanya.
    - nama_font: Nama font yang akan digunakan (default: 'Times New Roman').
    - ukuran_font: Ukuran font yang akan digunakan (default: 12).
    - tengah_align: Jika True, maka paragraf akan diatur menjadi rata tengah (default: False).
    - justify_align: Jika True, maka paragraf akan diatur menjadi rata kanan dan kiri (default: False).
    """
    sel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraf in sel.paragraphs:
        # Atur alignment
        if tengah_align:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif justify_align:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Atur font
        for run in paragraf.runs:
            run.font.name = nama_font
            run.font.size = Pt(ukuran_font)

# Pastikan juga untuk mengimpor WD_ALIGN_PARAGRAPH dari docx.enum.text

def atur_tinggi_baris(tabel, tinggi_cm):
    for baris in tabel.rows:
        for sel in baris.cells:
            sel.height = Cm(tinggi_cm)

from docx.shared import Cm

# Fungsi untuk mengatur tinggi baris dalam tabel
def atur_tinggi_baris(tabel, baris, tinggi_cm):
    for sel in tabel.rows[baris].cells:
        sel.height = Cm(tinggi_cm)



# raport by kelas


@app.route('/generate-rapor-rplxi/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_rplxi(id_kelas, id_semester):
    # Ambil semua siswa dalam kelas
    kelas = Kelas.query.get(id_kelas)
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)

    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
            # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()

            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,
                "Matematika": 9,
                "Bahasa Inggris": 10,
                "Rekayasa Perangkat Lunak": 11,
                "Projek Kreatif dan Kewirausahaan": 12,
                "Digital Marketing": 13,
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_newxi_rpl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)     

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan, justify_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter, tinggi_cm)

                counter += 1    

            # Isi kegiatan industri
            counter = 1  
            for kegiatan in kegiatan_industri:
                baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
                sel_nomor = baris.cells[0]  # Sel nomor
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_mitra_industri = baris.cells[1]  # Sel mitra industri
                sel_mitra_industri.text = kegiatan.mitra_induka
                atur_gaya_sel(sel_mitra_industri)

                sel_lokasi = baris.cells[2]  # Sel lokasi
                sel_lokasi.text = kegiatan.lokasi
                atur_gaya_sel(sel_lokasi, tengah_align=True)

                sel_keterangan = baris.cells[3]  # Sel keterangan
                sel_keterangan.text = kegiatan.keterangan
                atur_gaya_sel(sel_keterangan, tengah_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[2], counter, tinggi_cm)

                counter += 1

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[3].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[3].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[3].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            
            # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')



@app.route('/generate-rapor-tsmxi/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_tsmxi(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)

    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
    # ... sisa kode
           # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,

                "Matematika": 9,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 10,
                "Teknik Sepeda Motor": 11,
                "Projek Kreatif dan Kewirausahaan": 12,
                "Digital Marketing": 13,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_newxi_tsm.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan, justify_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter, tinggi_cm)

                counter += 1    

            # Isi kegiatan industri
            counter = 1  
            for kegiatan in kegiatan_industri:
                baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
                sel_nomor = baris.cells[0]  # Sel nomor
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_mitra_industri = baris.cells[1]  # Sel mitra industri
                sel_mitra_industri.text = kegiatan.mitra_induka
                atur_gaya_sel(sel_mitra_industri)

                sel_lokasi = baris.cells[2]  # Sel lokasi
                sel_lokasi.text = kegiatan.lokasi
                atur_gaya_sel(sel_lokasi, tengah_align=True)

                sel_keterangan = baris.cells[3]  # Sel keterangan
                sel_keterangan.text = kegiatan.keterangan
                atur_gaya_sel(sel_keterangan, tengah_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[2], counter, tinggi_cm)

                counter += 1

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[3].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[3].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[3].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            
            # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

    #return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{id_kelas}.zip')

@app.route('/generate-rapor-htlxi/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_htlxi(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporanf
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,

                "Matematika": 9,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 10,
                "Perhotelan": 11,
                "Projek Kreatif dan Kewirausahaan": 12,
                "Digital Marketing": 13,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_newxi_htl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)
            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan, justify_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter, tinggi_cm)

                counter += 1    

            # Isi kegiatan industri
            counter = 1  
            for kegiatan in kegiatan_industri:
                baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
                sel_nomor = baris.cells[0]  # Sel nomor
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_mitra_industri = baris.cells[1]  # Sel mitra industri
                sel_mitra_industri.text = kegiatan.mitra_induka
                atur_gaya_sel(sel_mitra_industri)

                sel_lokasi = baris.cells[2]  # Sel lokasi
                sel_lokasi.text = kegiatan.lokasi
                atur_gaya_sel(sel_lokasi, tengah_align=True)

                sel_keterangan = baris.cells[3]  # Sel keterangan
                sel_keterangan.text = kegiatan.keterangan
                atur_gaya_sel(sel_keterangan, tengah_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[2], counter, tinggi_cm)

                counter += 1

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[3].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[3].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[3].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            
            # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')



@app.route('/generate-rapor-rplx/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_rplx(id_kelas, id_semester):
    # Ambil semua siswa dalam kelas
    kelas = Kelas.query.get(id_kelas)
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()

    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Seni Musik": 7,
                "Mandarin": 8,

                "Matematika": 10,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 11,
                "Informatika": 12,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
                "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 14,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_x_rpl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)

            counter += 1    
            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')


@app.route('/generate-rapor-tsmx/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_tsmx(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()

    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Seni Musik": 7,
                "Mandarin": 8,

                "Matematika": 10,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 11,
                "Informatika": 12,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
                "Dasar-Dasar Otomotif": 14,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_x_tsm.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)

            counter += 1    
            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')


@app.route('/generate-rapor-htlx/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_htlx(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa,id_semester=id_semester).first()

    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Seni Musik": 7,
                "Mandarin": 8,

                "Matematika": 10,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 11,
                "Informatika": 12,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
                "Dasar-Dasar Perhotelan": 14,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_x_htl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)

            counter += 1    
            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

# end ..







def generate_excel(kelas_id):
    kelas = Kelas.query.get(kelas_id)
    mapel = Mapel.query.filter_by(id_kelas=kelas.id_kelas).all()

    # Membuat DataFrame
    data = {
        "Nama Mata Pelajaran": [m.nama_mapel for m in mapel],
        "No Id Mata Pelajaran": [m.id_mapel for m in mapel],
        "Guru Pengajar": [m.guru.nama_guru for m in mapel],
        "No Id Guru": [m.id_guru for m in mapel]
    }
    df = pd.DataFrame(data)

    # Mengonversi DataFrame ke Excel
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name="Laporan Kelas", index=False)

    excel_buffer.seek(0)
    return excel_buffer

@app.route('/download-laporan-kelas/<int:id_kelas>')
@login_required
def download_laporan_kelas(id_kelas):
    excel_buffer = generate_excel(id_kelas)
    return send_file(
        excel_buffer,
        as_attachment=True,
        download_name=f'laporan_kelas_{id_kelas}.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@app.route('/laporan_guru', methods=['GET'])
def laporan_gurublm():
    # Query data dan urutkan berdasarkan nama guru dan nama kelas
    data = db.session.query(Guru, Mapel, Kelas, Siswa)\
        .join(Mapel, Guru.id_guru == Mapel.id_guru)\
        .join(Kelas, Mapel.id_kelas == Kelas.id_kelas)\
        .join(Siswa, Kelas.id_kelas == Siswa.id_kelas)\
        .outerjoin(NilaiAkhir, (NilaiAkhir.id_guru == Guru.id_guru) & (NilaiAkhir.id_mapel == Mapel.id_mapel) & (NilaiAkhir.id_siswa == Siswa.id_siswa))\
        .filter(NilaiAkhir.id_nilai == None)\
        .order_by(Guru.nama_guru, Kelas.nama_kelas)\
        .all()

    # Membuat workbook baru
    wb = Workbook()
    ws = wb.active

    # Menambahkan header
    ws.append(["ID Guru", "Nama Guru", "Mata Pelajaran","Nama", "Kelas"])

    # Menambahkan data guru ke dalam worksheet
    for guru, mapel, kelas, siswa in data:
        ws.append([guru.id_guru, guru.nama_guru, mapel.nama_mapel,siswa.nama, kelas.nama_kelas ])

    # Menyimpan workbook ke dalam memory
    mem = BytesIO()
    wb.save(mem)
    mem.seek(0)

    # Mengirim file sebagai response
    return send_file(mem, as_attachment=True, download_name='laporan_guru.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/laporan_guru_sudah', methods=['GET'])
def laporan_gurusdh():
    from openpyxl.styles import Font, Alignment, PatternFill
    from collections import defaultdict
    from openpyxl.utils import get_column_letter

    # Query data guru, mapel, kelas, dan siswa
    data = db.session.query(Guru, Mapel, Kelas, Siswa)\
        .join(Mapel, Guru.id_guru == Mapel.id_guru)\
        .join(Kelas, Mapel.id_kelas == Kelas.id_kelas)\
        .join(Siswa, Kelas.id_kelas == Siswa.id_kelas)\
        .outerjoin(NilaiAkhir, (NilaiAkhir.id_guru == Guru.id_guru) & (NilaiAkhir.id_mapel == Mapel.id_mapel) & (NilaiAkhir.id_siswa == Siswa.id_siswa))\
        .order_by(Guru.nama_guru, Kelas.nama_kelas)\
        .all()

    # Mengelompokkan data berdasarkan guru
    grouped_data = defaultdict(list)
    for guru, mapel, kelas, siswa in data:
        grouped_data[guru.nama_guru].append((guru, mapel, kelas, siswa))

    # Membuat workbook baru
    wb = Workbook()
    del wb['Sheet']  # Hapus worksheet default yang dibuat saat inisialisasi Workbook
    
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")

    # Menambahkan data guru ke dalam worksheet masing-masing
    for guru_name, guru_data in grouped_data.items():
        ws = wb.create_sheet(title=guru_name[:31])  # Membuat sheet baru dengan nama guru, max 31 karakter
        
        # Menambahkan keterangan di sel A1
        ws['A1'] = f"Laporan Nama {guru_name}"
        ws['A1'].font = Font(bold=True, size=14)
        ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
        
        # Menggabungkan sel A1 dengan kolom terakhir dari header
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
        
        ws.append(["ID Guru", "Nama Guru", "Mata Pelajaran", "Nama Siswa", "Kelas", "Nilai Akhir", "Capaian Kompetensi"])  # Header
        
        # Set alignment for merged cells
        for cell in ws[1]:
            cell.alignment = Alignment(horizontal="center", vertical="center")

        for cell in ws[2]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        for guru, mapel, kelas, siswa in guru_data:
            nilaiakhir = 'Belum Ada'
            capaian_kompetensi = 'Belum Ada'
            nilaiakhir_obj = next((n for n in guru.nilai_akhir if n.id_mapel == mapel.id_mapel and n.id_siswa == siswa.id_siswa), None)
            if nilaiakhir_obj:
                nilaiakhir = nilaiakhir_obj.nilai
                capaian_kompetensi = nilaiakhir_obj.capaian_kompetensi  # Mengambil capaian kompetensi
            ws.append([guru.id_guru, guru.nama_guru, mapel.nama_mapel, siswa.nama, kelas.nama_kelas, nilaiakhir, capaian_kompetensi])

        # Mengatur lebar kolom secara otomatis
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if cell.coordinate not in ws.merged_cells and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column_letter].width = adjusted_width

    # Menyimpan workbook ke dalam memory
    mem = BytesIO()
    wb.save(mem)
    mem.seek(0)

    # Mengirim file sebagai response
    return send_file(mem, as_attachment=True, download_name='laporan_guru.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/peringkat', methods=['GET', 'POST'])
def peringkat_kelas():
    if request.method == 'POST':
        # Mendapatkan ID kelas yang dipilih dari formulir
        kelas_id = request.form.get('kelas_id')
        # Ambil data kelas berdasarkan ID kelas yang dipilih
        kelas = Kelas.query.get(kelas_id)

        # Ambil semua siswa di kelas tertentu
        siswa_kelas = Siswa.query.join(Kelas).filter(Kelas.id_kelas == kelas.id_kelas).all()

        data = []
        for s in siswa_kelas:
            # Hitung nilai rata-rata setiap siswa
            total_nilai = 0
            total_mapel = 0
            for n in s.nilai_akhir:
                total_nilai += n.nilai
                total_mapel += 1
            rata_rata = total_nilai / total_mapel

            data.append({
                "nama": s.nama,
                "rata_rata": rata_rata,
                "nis": s.nisn  # Tambahkan NIS ke dalam data siswa
            })

        # Urutkan data berdasarkan nilai rata-rata descending
        data = sorted(data, key=lambda x: x['rata_rata'], reverse=True)

        return render_template('peringkat_kelas.html',
                               daftar_kelas=Kelas.query.all(),  # Kirim daftar kelas ke template
                               nama_kelas=kelas.nama_kelas,
                               siswa=data,  # Kirim data yang sudah diperbarui ke template
                               peringkat1=data[0],
                               peringkat2=data[1],
                               peringkat3=data[2])

    # Jika metode GET, tampilkan halaman pertama kali dengan dropdown kosong
    return render_template('peringkat_kelas.html', daftar_kelas=Kelas.query.all())

@app.route('/absensi_harian', methods=['GET', 'POST'])
def absensi_harian():
    form = AbsensiHarianForm()
    form.id_kelas.choices = [(kelas.id_kelas, kelas.nama_kelas) for kelas in Kelas.query.order_by(Kelas.nama_kelas).all()]
    form.id_siswa.choices = [(siswa.id_siswa, siswa.nama) for siswa in Siswa.query.order_by(Siswa.nama).all()]

    if form.validate_on_submit():
        absensi = AbsensiHarian(
            id_siswa=form.id_siswa.data,
            tanggal=form.tanggal.data,
            status_kehadiran=form.status_kehadiran.data,
            keterangan=form.keterangan.data
        )
        db.session.add(absensi)
        db.session.commit()
        flash('Absensi berhasil ditambahkan')

        if form.status_kehadiran.data == 'alpha':
            siswa = Siswa.query.get(form.id_siswa.data)
            if siswa and siswa.orang_tua:
                nomor_telepon = siswa.orang_tua.nomor_whatsapp
                if nomor_telepon:
                    pesan = f"Siswa dengan nama {siswa.nama} tidak hadir tanpa keterangan (Alpha) pada tanggal {form.tanggal.data.strftime('%Y-%m-%d')}."
                    pesan_encoded = quote_plus(pesan)
                    whatsapp_url = f"https://api.whatsapp.com/send/?phone={nomor_telepon}&text={pesan_encoded}"
                    return redirect(whatsapp_url)
                else:
                    flash('Tidak ada nomor telepon orang tua yang terdaftar.')
            else:
                flash('Data orang tua tidak ditemukan.')

        return redirect(url_for('absensi_harian'))
    
    return render_template('absensi_harian.html', title='Absensi Harian', form=form)

"""
@app.route('/get_siswa_by_kelas/<int:kelas_id>', methods=['GET'])
def get_siswa_by_kelas(kelas_id):
    siswa = Siswa.query.filter_by(id_kelas=kelas_id).all()
    return jsonify([{'id_siswa': s.id_siswa, 'nama': s.nama} for s in siswa])
"""

@app.route('/get_orang_tua_by_siswa/<int:siswa_id>', methods=['GET'])
def get_orang_tua_by_siswa(siswa_id):
    orang_tua = OrangTua.query.filter_by(id_siswa=siswa_id).first()
    if orang_tua:
        return jsonify({'id_orang_tua': orang_tua.id_orang_tua, 'nama_orang_tua': orang_tua.nama_orang_tua})
    return jsonify({})

def kirim_pesan_bulk(tanggal_pilihan):
    tanggal = datetime.strptime(tanggal_pilihan, '%Y-%m-%d').date()  # Konversi string ke objek tanggal
    # Sisa kode tetap sama

    absensi_alpha = AbsensiHarian.query.filter_by(tanggal=tanggal, status_kehadiran='alpha').all()

    for absensi in absensi_alpha:
        siswa = Siswa.query.get(absensi.id_siswa)
        if siswa and siswa.orang_tua and siswa.orang_tua.nomor_whatsapp:
            nomor_telepon = siswa.orang_tua.nomor_whatsapp
            pesan = f"Siswa dengan nama {siswa.nama} tidak hadir tanpa keterangan pada tanggal {tanggal}."
            pesan_encoded = quote_plus(pesan)
            whatsapp_url = f"https://api.whatsapp.com/send/?phone={nomor_telepon}&text={pesan_encoded}"
            # Kirim pesan menggunakan requests atau metode lain
            requests.get(whatsapp_url)



@app.route('/kirim_pesan_bulk', methods=['POST'])
@login_required  # Jika diperlukan
def route_kirim_pesan_bulk():
    try:
        tanggal_pilihan = request.form['tanggal']
        kirim_pesan_bulk(tanggal_pilihan)
        flash('Pesan berhasil dikirim.', 'success')
    except Exception as e:
        flash(f'Gagal mengirim pesan: {e}', 'danger')
    return redirect(url_for('dashboard'))  # Atau route yang relevan

#17052024

# form rapor persiswa
@app.route('/form-generate-rapor-name')
@login_required
def form_generate_rapor():
    return render_template('form_generate_rapor.html')


# Generate by name
@app.route('/generate-rapor-nama')
@login_required
def generate_rapor_nama():
    nama_siswa = request.args.get('nama_siswa')
    # Retrieve student data and grades from the database based on the student's name
   
    siswa = Siswa.query.filter_by(nama=nama_siswa).first()
    
    # Check if the student was found
    if siswa is None:
        return "Siswa tidak ditemukan", 404
    id_siswa = siswa.id_siswa
    id_kelas = siswa.id_kelas
    nama_kelas = siswa.kelas.nama_kelas

    # Determine which report generation function to call based on the student's class
    if nama_kelas == "XI Rekayasa Perangkat Lunak":
        #return generate_rapor_kelas_rplxi(id_kelas)
        return generate_rapor_rplxi(id_siswa)
    elif nama_kelas == "XI Teknik Sepeda Motor":
        #return generate_rapor_kelas_tsmxi(id_kelas)
        return generate_rapor_tsmxi(id_siswa)
    elif nama_kelas == "XI Perhotelan":
        #return generate_rapor_kelas_htlxi(id_kelas)
        return generate_rapor_htlxi(id_siswa)
    elif nama_kelas == "X Rekayasa Perangkat Lunak":
        #return generate_rapor_kelas_rplx(id_kelas)
        return generate_rapor_rplx(id_siswa)
    elif nama_kelas == "X Teknik Sepeda Motor":
        #return generate_rapor_kelas_tsmx(id_kelas)
        return generate_rapor_tsmx(id_siswa)
    elif nama_kelas == "X Perhotelan":
        #return generate_rapor_kelas_htlx(id_kelas)
        return generate_rapor_htlx(id_siswa)
    else:
        return "Kelas tidak dikenal", 400


def generate_rapor_rplxi(id_siswa):
    mapel_ke_baris_rplxi = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Mandarin": 7,
        "Matematika": 9,
        "Bahasa Inggris": 10,
        "Rekayasa Perangkat Lunak": 11,
        "Projek Kreatif dan Kewirausahaan": 12,
        "Digital Marketing": 13,
    }
    return generate_rapor(id_siswa, 'template_rapor_newxi_rpl.docx', mapel_ke_baris_rplxi, include_kegiatan_industri=True)

def generate_rapor_tsmxi(id_siswa):
    mapel_ke_baris_tsmxi = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Mandarin": 7,
        "Matematika": 9,
        "Bahasa Inggris": 10,
        "Teknik Sepeda Motor": 11,
        "Projek Kreatif dan Kewirausahaan": 12,
        "Digital Marketing": 13,
    }
    return generate_rapor(id_siswa, 'template_rapor_newxi_tsm.docx', mapel_ke_baris_tsmxi, include_kegiatan_industri=True)

def generate_rapor_htlxi(id_siswa):
    mapel_ke_baris_htlxi = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Mandarin": 7,
        "Matematika": 9,
        "Bahasa Inggris": 10,
        "Perhotelan": 11,
        "Projek Kreatif dan Kewirausahaan": 12,
        "Digital Marketing": 13,
    }
    return generate_rapor(id_siswa, 'template_rapor_newxi_htl.docx', mapel_ke_baris_htlxi, include_kegiatan_industri=True)

def generate_rapor_rplx(id_siswa):
    mapel_ke_baris_rplx = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Seni Musik": 7,
        "Mandarin": 8,
        "Matematika": 10,  
        "Bahasa Inggris": 11,
        "Informatika": 12,
        "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
        "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 14,
    }
    return generate_rapor(id_siswa, 'template_rapor_x_rpl.docx', mapel_ke_baris_rplx)

def generate_rapor_tsmx(id_siswa):
    mapel_ke_baris_tsmx = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Seni Musik": 7,
        "Mandarin": 8,
        "Matematika": 10,  
        "Bahasa Inggris": 11,
        "Informatika": 12,
        "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
        "Dasar-Dasar Otomotif": 14,
    }
    return generate_rapor(id_siswa, 'template_rapor_x_tsm.docx', mapel_ke_baris_tsmx)

def generate_rapor_htlx(id_siswa):
    mapel_ke_baris_htlx = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Seni Musik": 7,
        "Mandarin": 8,
        "Matematika": 10,  
        "Bahasa Inggris": 11,
        "Informatika": 12,
        "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
        "Dasar-Dasar Perhotelan": 14,
    }
    return generate_rapor(id_siswa, 'template_rapor_x_htl.docx', mapel_ke_baris_htlx)

# end form rapor persiswa


# end form rapor persiswa

def generate_rapor(id_siswa, template_path, mapel_ke_baris, include_kegiatan_industri=False):
    # Ambil data siswa dan nilai dari database
    siswa = Siswa.query.get(id_siswa)
    nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=id_siswa).all()
    ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=id_siswa).all()
    rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=id_siswa).first()
   
    # Periksa jika kegiatan_industri perlu dimasukkan berdasarkan template
    if "rplxi" in template_path or "tsmxi" in template_path or "htlxi" in template_path or include_kegiatan_industri:
        kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=id_siswa).all()
    else:
        kegiatan_industri = [] 

    # Ambil informasi semester dari salah satu nilai akhir
    if nilai_akhir:
        semester = Semester.query.get(nilai_akhir[0].id_semester)
    else:
        semester = None

    # Buka template
    doc = Document(template_path)
    
    # Isi bagian data siswa
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '{Nama}' in run.text and siswa.nama is not None:
                run.text = run.text.replace('{Nama}', siswa.nama)
            if '{Kelas}' in run.text:
                if siswa.kelas.nama_kelas is not None:
                    run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                else:
                    run.text = run.text.replace('{Kelas}', 'Default Class')
            if '{NISN/NIS}' in run.text and siswa.nisn is not None:
                run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
            if semester:
                if '{Semester}' in run.text:
                    run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                if '{TahunAjaran}' in run.text:
                    run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

    # Isi nilai akhir
    for nilai in nilai_akhir:
        baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
        if baris is not None:
            # Sel nama mata pelajaran
            sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
            sel_nama_mapel.text = nilai.mapel.nama_mapel
            atur_gaya_sel(sel_nama_mapel)

            # Sel nilai
            sel_nilai = doc.tables[0].rows[baris].cells[2]
            sel_nilai.text = str(nilai.nilai)
            atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

            # Sel capaian kompetensi
            sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
            sel_capaian_kompetensi.text = nilai.capaian_kompetensi
            atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)

    # Isi kegiatan ekstrakurikuler
    counter = 1
    for ekstra in ekstrakurikuler:
        baris = doc.tables[1].add_row()
        sel_nomor = baris.cells[0]
        sel_nomor.text = str(counter)
        atur_gaya_sel(sel_nomor, tengah_align=True)

        sel_kegiatan = baris.cells[1]
        sel_kegiatan.text = ekstra.kegiatan
        atur_gaya_sel(sel_kegiatan)

        sel_keterangan = baris.cells[2]
        sel_keterangan.text = ekstra.keterangan
        atur_gaya_sel(sel_keterangan)

        counter += 1    

    # Isi kegiatan industri jika ada
    counter = 1
    if kegiatan_industri:
        for kegiatan in kegiatan_industri:
            baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
            sel_nomor = baris.cells[0]  # Sel nomor
            sel_nomor.text = str(counter)
            atur_gaya_sel(sel_nomor, tengah_align=True)

            sel_mitra_industri = baris.cells[1]  # Sel mitra industri
            sel_mitra_industri.text = kegiatan.mitra_induka
            atur_gaya_sel(sel_mitra_industri)

            sel_lokasi = baris.cells[2]  # Sel lokasi
            sel_lokasi.text = kegiatan.lokasi
            atur_gaya_sel(sel_lokasi, tengah_align=True)

            sel_keterangan = baris.cells[3]  # Sel keterangan
            sel_keterangan.text = kegiatan.keterangan
            atur_gaya_sel(sel_keterangan, tengah_align=True)

            counter += 1    

    # Isi rekapitulasi kehadiran
    # Periksa jumlah tabel dalam dokumen
    table_index = 3 if kegiatan_industri else 2
    if rekap_absensi:
        # Sel total sakit
        sel_total_sakit = doc.tables[table_index].rows[0].cells[1]
        sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
        atur_gaya_sel(sel_total_sakit)

        # Sel total izin
        sel_total_izin = doc.tables[table_index].rows[1].cells[1]
        sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
        atur_gaya_sel(sel_total_izin)

        # Sel total tanpa keterangan
        sel_total_tanpa_keterangan = doc.tables[table_index].rows[2].cells[1]
        sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
        atur_gaya_sel(sel_total_tanpa_keterangan)

    # Ambil bagian atas halaman
    section = doc.sections[0]

    # Memanggil footer
    footer = section.footer

    # Tambahkan teks footer dengan informasi siswa
    footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
    footer_para = footer.paragraphs[0]
    footer_para.text = footer_text

    # Simpan dokumen ke dalam buffer dan kirim sebagai file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'rapor_{siswa.nama}.docx')


@app.route('/generate-raporsiswa', methods=['GET', 'POST'])
@login_required
def generate_raporsiswa():
    kelas_list = Kelas.query.all()
    siswa_list = Siswa.query.all()
    semester_list = Semester.query.all()

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        siswa_id = request.form.get('siswa')
        semester_id = request.form.get('semester')

        if kelas_id and siswa_id and semester_id:
            siswa = Siswa.query.filter_by(id_siswa=siswa_id).first()
            nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            semester = Semester.query.filter_by(id=semester_id).first()
            
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,
                "Seni Musik": 7,
                "Matematika": 8,
                "Bahasa Inggris": 9,
                "Rekayasa Perangkat Lunak": 10,
                "Teknik Sepeda Motor": 10,
                "Perhotelan": 10,
                "Projek Kreatif dan Kewirausahaan": 11,
                "Informatika": 11,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 12,
                "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 13,
                "Digital Marketing": 12,
                "Dasar-Dasar Otomotif": 13,
                "Dasar-Dasar Perhotelan": 13,
            }

            return render_template('rapor_siswa.html', siswa=siswa, nilai_akhir=nilai_akhir, ekstrakurikuler=ekstrakurikuler, rekap_absensi=rekap_absensi, kegiatan_industri=kegiatan_industri, mapel_ke_baris=mapel_ke_baris, kelas_list=kelas_list, siswa_list=siswa_list, semester_list=semester_list, semester=semester)

    return render_template('rapor_siswa.html', kelas_list=kelas_list, siswa_list=siswa_list, semester_list=semester_list)




@app.route('/get_siswa_by_kelas/<kelas_id>')
@login_required
def get_siswa_by_kelas(kelas_id):
    siswa_list = Siswa.query.filter_by(id_kelas=kelas_id).all()
    siswa_data = [{'id_siswa': s.id_siswa, 'nama': s.nama} for s in siswa_list]
    return jsonify(siswa_data)

@app.route('/generate-rapor-kelas/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas(id_kelas, id_semester):
    # Retrieve class data from the database
    kelas = Kelas.query.get(id_kelas)
    
    # Check if the class was found
    if kelas is None:
        return "Kelas tidak ditemukan", 404
    
    nama_kelas = kelas.nama_kelas

    # Determine which report generation function to call based on the class and semester
    if nama_kelas == "XI Rekayasa Perangkat Lunak":
        return generate_rapor_kelas_rplxi(id_kelas, id_semester)
    elif nama_kelas == "XI Teknik Sepeda Motor":
        return generate_rapor_kelas_tsmxi(id_kelas, id_semester)
    elif nama_kelas == "XI Perhotelan":
        return generate_rapor_kelas_htlxi(id_kelas, id_semester)
    elif nama_kelas == "X Rekayasa Perangkat Lunak":
        return generate_rapor_kelas_rplx(id_kelas, id_semester)
    elif nama_kelas == "X Teknik Sepeda Motor":
        return generate_rapor_kelas_tsmx(id_kelas, id_semester)
    elif nama_kelas == "X Perhotelan":
        return generate_rapor_kelas_htlx(id_kelas, id_semester)
    else:
        return "Kelas tidak dikenal", 400


@app.route('/download/rekap_absensi')
def download_rekap_absensi():
    # Path to the Excel file for rekap absensi
    excel_file_path = 'download/rekap_absensi.xlsx'
    return send_file(excel_file_path, as_attachment=True)

@app.route('/download/rekap_kegiatan_industri')
def download_rekap_kegiatan_industri():
    # Path to the Excel file for rekap kegiatan industri
    excel_file_path = 'download/rekap_kegiatan_industri.xlsx'
    return send_file(excel_file_path, as_attachment=True)

@app.route('/download/rekap_kegiatan_ekstrakurikuler')
def download_rekap_kegiatan_ekstrakurikuler():
    # Path to the Excel file for rekap kegiatan ekstrakurikuler
    excel_file_path = 'download/rekap_kegiatan_ekstrakurikuler.xlsx'
    return send_file(excel_file_path, as_attachment=True) 



# routes.py
# routes.py
@app.route('/api/student_count_per_class')
def get_student_count_per_class():
    results = db.session.query(Kelas.nama_kelas, func.count(Siswa.id_siswa)).join(Siswa).group_by(Kelas.nama_kelas).all()
    data = [{'kelas': row[0], 'jumlah': row[1]} for row in results]
    return jsonify(data)

@app.route('/api/average_score_per_subject')
def get_average_score_per_subject():
    results = db.session.query(Mapel.nama_mapel, func.avg(NilaiAkhir.nilai)).join(NilaiAkhir).group_by(Mapel.nama_mapel).all()
    data = [{'mapel': row[0], 'rata_rata': row[1]} for row in results]
    return jsonify(data)

@app.route('/api/attendance_per_class')
def get_attendance_per_class():
    results = db.session.query(
        Kelas.nama_kelas,
        func.avg(RekapAbsensi.total_sakit).label('avg_sakit'),
        func.avg(RekapAbsensi.total_izin).label('avg_izin'),
        func.avg(RekapAbsensi.total_tanpa_keterangan).label('avg_tanpa_keterangan')
    ).join(Siswa, Siswa.id_kelas == Kelas.id_kelas).join(RekapAbsensi, RekapAbsensi.id_siswa == Siswa.id_siswa).group_by(Kelas.nama_kelas).all()

    data = [{'kelas': row.nama_kelas, 'sakit': row.avg_sakit, 'izin': row.avg_izin, 'alfa': row.avg_tanpa_keterangan} for row in results]
    return jsonify(data)



@app.route('/api/top_students')
def get_top_students():
    results = db.session.query(
        Siswa.nama,
        func.avg(NilaiAkhir.nilai).label('avg_nilai')
    ).join(NilaiAkhir).group_by(Siswa.id_siswa).order_by(func.avg(NilaiAkhir.nilai).desc()).limit(3).all()

    data = [{'nama': row.nama, 'avg_nilai': row.avg_nilai} for row in results]
    return jsonify(data)


@app.route('/api/semesters')
def get_semesters():
    semesters = Semester.query.filter(Semester.semester.in_(['II', 'IV'])).all()
    data = [{'id': semester.id, 'tahun_ajaran': semester.tahun_ajaran, 'semester': semester.semester} for semester in semesters]
    return jsonify(data)


@app.route('/api/classes')
def get_classes():
    classes = Kelas.query.order_by(Kelas.nama_kelas.asc()).all()
    data = [{'id': kelas.id_kelas, 'nama': kelas.nama_kelas} for kelas in classes]
    return jsonify(data)

@app.route('/api/progress_per_class/<int:class_id>/<int:semester_id>')
def get_progress_per_class(class_id, semester_id):
    # Calculate the total number of students in the class
    total_students = db.session.query(Siswa).filter_by(id_kelas=class_id).count()

    results = db.session.query(
        Guru.nama_guru,
        Mapel.nama_mapel,
        func.count(NilaiAkhir.id_nilai).label('nilai_count'),
        (func.count(NilaiAkhir.id_nilai) * 100 / total_students).label('progress')
    ).join(NilaiAkhir, (NilaiAkhir.id_guru == Guru.id_guru) & (NilaiAkhir.id_kelas == class_id) & (NilaiAkhir.id_semester == semester_id)).join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel).group_by(Guru.nama_guru, Mapel.nama_mapel).all()

    data = [{'guru': row.nama_guru, 'mapel': row.nama_mapel, 'progress': row.progress} for row in results]
    return jsonify(data)

@app.route('/api/ranking_per_class/<int:class_id>/<int:semester_id>')
def get_ranking_per_class(class_id, semester_id):
    results = db.session.query(
        Siswa.nama,
        func.avg(NilaiAkhir.nilai).label('avg_nilai')
    ).join(NilaiAkhir).filter(
        Siswa.id_kelas == class_id,
        NilaiAkhir.id_semester == semester_id
    ).group_by(Siswa.id_siswa).order_by(func.avg(NilaiAkhir.nilai).desc()).all()

    data = [{'nama': row.nama, 'avg_nilai': round(row.avg_nilai, 2)} for row in results]
    return jsonify(data)

@app.route('/api/semesterall')
def get_semester_all():
    semesters = Semester.query.order_by(Semester.tahun_ajaran.asc(), Semester.semester.asc()).all()
    data = [{'id': semester.id, 'tahun_ajaran': semester.tahun_ajaran, 'semester': semester.semester} for semester in semesters]
    return jsonify(data)

@app.route('/api/class_absensi/<int:class_id>/<int:semester_id>')
def class_absensi(class_id, semester_id):
    results = db.session.query(
        Siswa.nama,
        func.sum(RekapAbsensi.total_sakit).label('total_sakit'),
        func.sum(RekapAbsensi.total_izin).label('total_izin'),
        func.sum(RekapAbsensi.total_tanpa_keterangan).label('total_tanpa_keterangan')
    ).join(RekapAbsensi).filter(
        Siswa.id_kelas == class_id,
        RekapAbsensi.id_semester == semester_id
    ).group_by(Siswa.id_siswa).all()

    data = [{'nama': row.nama, 'total_sakit': row.total_sakit, 'total_izin': row.total_izin, 'total_tanpa_keterangan': row.total_tanpa_keterangan} for row in results]
    return jsonify(data)




@app.route('/api/classes_absensi')
def get_classes_absensi():
    classes = Kelas.query.order_by(Kelas.nama_kelas.asc()).all()
    data = [{'id': kelas.id_kelas, 'nama': kelas.nama_kelas} for kelas in classes]
    return jsonify(data)

@app.route('/api/students_in_class/<int:class_id>')
def get_students_in_class(class_id):
    students = Siswa.query.filter_by(id_kelas=class_id).order_by(Siswa.nama.asc()).all()
    data = [{'id': student.id_siswa, 'nama': student.nama} for student in students]
    return jsonify(data)

@app.route('/api/student_subject_scores/<int:student_id>')
def get_student_subject_scores(student_id):
    results = db.session.query(
        Mapel.nama_mapel,
        Semester.semester,
        func.avg(NilaiAkhir.nilai).label('avg_nilai')
    ).join(NilaiAkhir, NilaiAkhir.id_mapel == Mapel.id_mapel)\
     .join(Semester, Semester.id == NilaiAkhir.id_semester)\
     .filter(NilaiAkhir.id_siswa == student_id)\
     .group_by(Mapel.nama_mapel, Semester.semester)\
     .order_by(Semester.semester).all()

    data = [{'mapel': row.nama_mapel, 'semester': row.semester, 'avg_nilai': round(row.avg_nilai, 2)} for row in results]
    return jsonify(data)


@app.route('/charts')
def charts():
    return render_template('charts.html')



@app.route('/progress_guru')
def progress_guru():
    return render_template('progress_guru.html')

@app.route('/api/performance_over_time/<int:student_id>')
def get_performance_over_time(student_id):
    results = db.session.query(
        Semester.semester,
        func.avg(NilaiAkhir.nilai).label('avg_nilai')
    ).join(NilaiAkhir, NilaiAkhir.id_semester == Semester.id) \
     .filter(NilaiAkhir.id_siswa == student_id) \
     .group_by(Semester.semester) \
     .order_by(Semester.semester).all()

    data = [{'semester': row.semester, 'avg_nilai': round(row.avg_nilai, 2)} for row in results]
    return jsonify(data)


@app.route('/api/subject_difficulty')
def get_subject_difficulty():
    results = db.session.query(
        Mapel.nama_mapel,
        func.avg(NilaiAkhir.nilai).label('avg_nilai'),
        func.count(NilaiAkhir.id_siswa).label('jumlah_siswa')
    ).group_by(Mapel.nama_mapel).all()

    data = [{'mapel': row.nama_mapel, 'avg_nilai': round(row.avg_nilai, 2), 'jumlah_siswa': row.jumlah_siswa} for row in results]
    return jsonify(data)


@app.route('/chart1')
def chart1():
    return render_template('chart1.html')   

@app.route('/api/top_performers/<int:class_id>/<int:semester_id>')
def get_top_performers(class_id, semester_id):
    top_students = db.session.query(
        Siswa.nama,
        Mapel.nama_mapel,
        func.avg(NilaiAkhir.nilai).label('avg_nilai')
    ).join(NilaiAkhir, Siswa.id_siswa == NilaiAkhir.id_siswa)\
    .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
    .filter(Siswa.id_kelas == class_id, NilaiAkhir.id_semester == semester_id)\
    .group_by(Siswa.id_siswa, Mapel.id_mapel)\
    .order_by(func.avg(NilaiAkhir.nilai).desc())\
    .limit(5).all()

    # Organizing the data for radar chart
    students = {}
    subjects = set()

    for student, subject, avg_nilai in top_students:
        if student not in students:
            students[student] = {}
        students[student][subject] = avg_nilai
        subjects.add(subject)

    subjects = list(subjects)
    data = {
        'subjects': subjects,
        'students': [
            {
                'name': student,
                'scores': [students[student].get(subject, 0) for subject in subjects]
            }
        for student in students]
    }

    return jsonify(data)


@app.route('/api/student_performance/<int:class_id>/<int:semester_id>/<int:student_id>')
@login_required
def get_student_performance(class_id, semester_id, student_id):
    # Dapatkan data siswa, nilai akhir, dan rata-rata kelas berdasarkan id kelas, id semester, dan id siswa
    nilai_akhir = db.session.query(NilaiAkhir, Mapel.nama_mapel).join(Mapel).filter(
        NilaiAkhir.id_siswa == student_id,
        NilaiAkhir.id_semester == semester_id
    ).all()

    if not nilai_akhir:
        return jsonify({"error": "No data found"}), 404

    # Dapatkan rata-rata nilai untuk setiap mata pelajaran di kelas tersebut
    class_average = db.session.query(
        Mapel.nama_mapel,
        func.avg(NilaiAkhir.nilai).label('avg_nilai')
    ).join(NilaiAkhir).join(Siswa).filter(
        Siswa.id_kelas == class_id,
        NilaiAkhir.id_semester == semester_id
    ).group_by(Mapel.nama_mapel).all()

    # Buat dictionary untuk rata-rata kelas
    class_avg_dict = {item.nama_mapel: item.avg_nilai for item in class_average}

    subjects = []
    student_scores = []
    class_averages = []

    # Isi data untuk response
    for nilai, mapel in nilai_akhir:
        subjects.append(mapel)
        student_scores.append(nilai.nilai)
        class_averages.append(class_avg_dict.get(mapel, 0))  # Jika tidak ada data rata-rata, gunakan 0

    data = {
        "subjects": subjects,
        "scores": student_scores,
        "class_average": class_averages
    }

    return jsonify(data)


@app.route('/api/subject_scores/<int:class_id>/<int:semester_id>/<int:student_id>', methods=['GET'])
def get_subject_scores(class_id, semester_id, student_id):
    subject_scores = db.session.query(
        Mapel.nama_mapel, 
        NilaiAkhir.nilai
    ).join(NilaiAkhir).filter(
        NilaiAkhir.id_siswa == student_id,
        NilaiAkhir.id_kelas == class_id,
        NilaiAkhir.id_semester == semester_id
    ).all()

    data = [{'mapel': mapel, 'nilai': nilai} for mapel, nilai in subject_scores]
    return jsonify(data)


@app.route('/api/attendance/<int:student_id>', methods=['GET'])
def get_attendance(student_id):
    attendance = db.session.query(RekapAbsensi).filter_by(id_siswa=student_id).first()

    if attendance:
        data = {
            'total_sakit': attendance.total_sakit,
            'total_izin': attendance.total_izin,
            'total_tanpa_keterangan': attendance.total_tanpa_keterangan
        }
    else:
        data = {
            'total_sakit': 0,
            'total_izin': 0,
            'total_tanpa_keterangan': 0
        }
    return jsonify(data)

@app.route('/api/ranking/<int:class_id>/<int:semester_id>', methods=['GET'])
def get_ranking(class_id, semester_id):
    ranking_data = db.session.query(
        Siswa.nama,
        func.rank().over(
            order_by=func.avg(NilaiAkhir.nilai).desc()
        ).label('rank')
    ).join(NilaiAkhir).filter(
        NilaiAkhir.id_kelas == class_id,
        NilaiAkhir.id_semester == semester_id
    ).group_by(Siswa.id_siswa).all()

    data = [{'nama': nama, 'rank': rank} for nama, rank in ranking_data]
    return jsonify(data)








class_map = {
    'XI Rekayasa Perangkat Lunak': {
        'template': 'template_rapor_newxi_rpl.docx',
        'mapel_ke_baris': {
            "Pendidikan Agama Islam dan Budi Pekerti": 2,
            "Pendidikan Agama Kristen dan Budi Pekerti": 2,
            "Pendidikan Agama Katolik dan Budi Pekerti": 2,
            "Pendidikan Pancasila": 3,
            "Bahasa Indonesia": 4,
            "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
            "Sejarah": 6,
            "Mandarin": 7,
            "Matematika": 9,
            "Bahasa Inggris": 10,
            "Rekayasa Perangkat Lunak": 11,
            "Projek Kreatif dan Kewirausahaan": 12,
            "Digital Marketing": 13,
        }
    },
    'XI Teknik Sepeda Motor': {
        'template': 'template_rapor_newxi_tsm.docx',
        'mapel_ke_baris': {
            "Pendidikan Agama Islam dan Budi Pekerti": 2,
            "Pendidikan Agama Kristen dan Budi Pekerti": 2,
            "Pendidikan Agama Katolik dan Budi Pekerti": 2,
            "Pendidikan Pancasila": 3,
            "Bahasa Indonesia": 4,
            "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
            "Sejarah": 6,
            "Mandarin": 7,
            "Matematika": 9,
            "Bahasa Inggris": 10,
            "Teknik Sepeda Motor": 11,
            "Projek Kreatif dan Kewirausahaan": 12,
            "Digital Marketing": 13,
        }
    },
    'XI Perhotelan': {
        'template': 'template_rapor_newxi_htl.docx',
        'mapel_ke_baris': {
            "Pendidikan Agama Islam dan Budi Pekerti": 2,
            "Pendidikan Agama Kristen dan Budi Pekerti": 2,
            "Pendidikan Agama Katolik dan Budi Pekerti": 2,
            "Pendidikan Pancasila": 3,
            "Bahasa Indonesia": 4,
            "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
            "Sejarah": 6,
            "Mandarin": 7,
            "Matematika": 9,
            "Bahasa Inggris": 10,
            "Perhotelan": 11,
            "Projek Kreatif dan Kewirausahaan": 12,
            "Digital Marketing": 13,
        }
    },
    'X Rekayasa Perangkat Lunak': {
        'template': 'template_rapor_x_rpl.docx',
        'mapel_ke_baris': {
            "Pendidikan Agama Islam dan Budi Pekerti": 2,
            "Pendidikan Agama Kristen dan Budi Pekerti": 2,
            "Pendidikan Agama Katolik dan Budi Pekerti": 2,
            "Pendidikan Pancasila": 3,
            "Bahasa Indonesia": 4,
            "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
            "Sejarah": 6,
            "Seni Musik": 7,
            "Mandarin": 8,
            "Matematika": 10,  
            "Bahasa Inggris": 11,
            "Informatika": 12,
            "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
            "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 14,
        }
    },
    'X Teknik Sepeda Motor': {
        'template': 'template_rapor_x_tsm.docx',
        'mapel_ke_baris': {
            "Pendidikan Agama Islam dan Budi Pekerti": 2,
            "Pendidikan Agama Kristen dan Budi Pekerti": 2,
            "Pendidikan Agama Katolik dan Budi Pekerti": 2,
            "Pendidikan Pancasila": 3,
            "Bahasa Indonesia": 4,
            "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
            "Sejarah": 6,
            "Seni Musik": 7,
            "Mandarin": 8,
            "Matematika": 10,  
            "Bahasa Inggris": 11,
            "Informatika": 12,
            "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
            "Dasar-Dasar Otomotif": 14,
        }
    },
    'X Perhotelan': {
        'template': 'template_rapor_x_htl.docx',
        'mapel_ke_baris': {
            "Pendidikan Agama Islam dan Budi Pekerti": 2,
            "Pendidikan Agama Kristen dan Budi Pekerti": 2,
            "Pendidikan Agama Katolik dan Budi Pekerti": 2,
            "Pendidikan Pancasila": 3,
            "Bahasa Indonesia": 4,
            "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
            "Sejarah": 6,
            "Seni Musik": 7,
            "Mandarin": 8,
            "Matematika": 10,  
            "Bahasa Inggris": 11,
            "Informatika": 12,
            "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
            "Dasar-Dasar Perhotelan": 14,
        }
    }
}

@app.route('/generate-rapor/<int:id_kelas>/<int:id_semester>/<int:id_siswa>')
@login_required
def generate_rapor(id_kelas, id_semester, id_siswa):
    # Ambil data siswa, kelas, dan semester
    kelas = Kelas.query.get(id_kelas)
    siswa = Siswa.query.get(id_siswa)
    semester = Semester.query.get(id_semester)

    # Periksa jika data siswa, kelas, atau semester tidak ditemukan
    if not siswa or not kelas or not semester:
        return "Data tidak ditemukan", 404

    # Pilih peta mata pelajaran dan template berdasarkan kelas
    class_info = class_map.get(kelas.nama_kelas)
    if not class_info:
        return "Template dan peta mata pelajaran untuk kelas ini tidak ditemukan", 404

    template_path = class_info['template']
    mapel_ke_baris = class_info['mapel_ke_baris']

    # Mengambil nilai akhir, ekstrakurikuler, rekap absensi, dan kegiatan industri
    nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=id_siswa, id_semester=id_semester).all()
    ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=id_siswa, id_semester=id_semester).all()
    rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=id_siswa, id_semester=id_semester).first()
    kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=id_siswa, id_semester=id_semester).all()

    # Buka template
    doc = Document(template_path)

    # Isi bagian data siswa
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '{Nama}' in run.text:
                run.text = run.text.replace('{Nama}', siswa.nama)
            if '{Kelas}' in run.text:
                run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
            if '{NISN/NIS}' in run.text:
                run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
            if '{Semester}' in run.text:
                run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
            if '{TahunAjaran}' in run.text:
                run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

    # Isi nilai akhir
    for nilai in nilai_akhir:
        baris = mapel_ke_baris.get(nilai.mapel.nama_mapel)
        if baris is not None:
            # Sel nama mata pelajaran
            sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
            sel_nama_mapel.text = nilai.mapel.nama_mapel
            atur_gaya_sel(sel_nama_mapel)

            # Sel nilai
            sel_nilai = doc.tables[0].rows[baris].cells[2]
            sel_nilai.text = str(nilai.nilai)
            atur_gaya_sel(sel_nilai, tengah_align=True)

            # Sel capaian kompetensi
            sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
            sel_capaian_kompetensi.text = nilai.capaian_kompetensi
            atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)

    # Isi kegiatan ekstrakurikuler
    counter = 1
    tinggi_cm = 0.8  # Tinggi baris dalam cm
    for ekstra in ekstrakurikuler:
        baris = doc.tables[1].add_row()
        sel_nomor = baris.cells[0]
        sel_nomor.text = str(counter)
        atur_gaya_sel(sel_nomor, tengah_align=True)

        sel_kegiatan = baris.cells[1]
        sel_kegiatan.text = ekstra.kegiatan
        atur_gaya_sel(sel_kegiatan)

        sel_keterangan = baris.cells[2]
        sel_keterangan.text = ekstra.keterangan
        atur_gaya_sel(sel_keterangan, justify_align=True)

        # Mengatur tinggi baris
        atur_tinggi_baris(doc.tables[1], counter, tinggi_cm)

        counter += 1    

    # Isi kegiatan industri jika tabel ada
    if "Mitra Induka" in [cell.text for cell in doc.tables[2].rows[0].cells]:
        counter = 1  
        for kegiatan in kegiatan_industri:
            baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
            sel_nomor = baris.cells[0]  # Sel nomor
            sel_nomor.text = str(counter)
            atur_gaya_sel(sel_nomor, tengah_align=True)

            sel_mitra_industri = baris.cells[1]  # Sel mitra industri
            sel_mitra_industri.text = kegiatan.mitra_induka
            atur_gaya_sel(sel_mitra_industri)

            sel_lokasi = baris.cells[2]  # Sel lokasi
            sel_lokasi.text = kegiatan.lokasi
            atur_gaya_sel(sel_lokasi, tengah_align=True)

            sel_keterangan = baris.cells[3]  # Sel keterangan
            sel_keterangan.text = kegiatan.keterangan
            atur_gaya_sel(sel_keterangan, tengah_align=True)

            # Mengatur tinggi baris
            atur_tinggi_baris(doc.tables[2], counter, tinggi_cm)

            counter += 1

    # Isi rekapitulasi kehadiran
    if rekap_absensi:
        table_index = 2 if len(doc.tables) == 3 else 3
        # Sel total sakit
        sel_total_sakit = doc.tables[table_index].rows[0].cells[1]
        sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
        atur_gaya_sel(sel_total_sakit)

        # Sel total izin
        sel_total_izin = doc.tables[table_index].rows[1].cells[1]
        sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
        atur_gaya_sel(sel_total_izin)

        # Sel total tanpa keterangan
        sel_total_tanpa_keterangan = doc.tables[table_index].rows[2].cells[1]
        sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
        atur_gaya_sel(sel_total_tanpa_keterangan)

    # Ambil bagian atas halaman
    section = doc.sections[0]

    # Memanggil footer
    footer = section.footer

    # Tambahkan teks footer dengan informasi siswa
    footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
    footer_para = footer.paragraphs[0]
    footer_para.text = footer_text

    # Simpan dokumen ke dalam buffer dan kirim sebagai file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'rapor_{siswa.nama}.docx')

@app.route('/generate-rapor-form')
@login_required
def generate_rapor_form():
    return render_template('generate_rapor_form.html')


@app.route('/download_files')
@login_required
def download_files():
    kelas = Kelas.query.order_by(Kelas.nama_kelas.asc()).all()
    return render_template('download_files.html', kelas=kelas)

@app.route('/bulk_delete_nilai_akhir_v2', methods=['GET', 'POST'])
@login_required
def bulk_delete_nilai_akhir_v2():
    kelas = Kelas.query.all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran).all()
    results = []

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        mapel_id = request.form.get('mapel')
        semester_id = request.form.get('semester')
        token = request.form.get('token')

        app.logger.info(f"Received POST with kelas_id={kelas_id}, guru_id={guru_id}, mapel_id={mapel_id}, semester_id={semester_id}, token={token}")

        if token == "123456" and kelas_id and guru_id and mapel_id and semester_id:
            NilaiAkhir.query.filter_by(id_kelas=kelas_id, id_guru=guru_id, id_mapel=mapel_id, id_semester=semester_id).delete()
            db.session.commit()
            flash('Data nilai akhir berhasil dihapus.', 'success')
        elif token != "123456":
            flash('Token salah. Data tidak dihapus.', 'danger')
        else:
            flash('Lengkapi semua pilihan.', 'danger')

    elif request.method == 'GET':
        kelas_id = request.args.get('kelas')
        guru_id = request.args.get('guru')
        mapel_id = request.args.get('mapel')
        semester_id = request.args.get('semester')

        app.logger.info(f"Received GET with kelas_id={kelas_id}, guru_id={guru_id}, mapel_id={mapel_id}, semester_id={semester_id}")

        if kelas_id and guru_id and mapel_id and semester_id:
            results = db.session.query(NilaiAkhir, Siswa, Mapel, Guru)\
                .join(Siswa, NilaiAkhir.id_siswa == Siswa.id_siswa)\
                .join(Mapel, NilaiAkhir.id_mapel == Mapel.id_mapel)\
                .join(Guru, NilaiAkhir.id_guru == Guru.id_guru)\
                .filter(NilaiAkhir.id_kelas == kelas_id, NilaiAkhir.id_guru == guru_id, NilaiAkhir.id_mapel == mapel_id, NilaiAkhir.id_semester == semester_id)\
                .all()

    return render_template('bulk_delete_nilai_akhir.html', kelas=kelas, semesters=semesters, results=results)


@app.route('/get-guru-by-kelas/<int:kelas_id>')
def get_guru_by_kelas33(kelas_id):
    guru_list = db.session.query(Guru).join(Mapel).filter(Mapel.id_kelas == kelas_id).distinct().all()
    return jsonify([{'id': guru.id_guru, 'name': guru.nama_guru} for guru in guru_list])

@app.route('/get-mapel-by-guru/<int:guru_id>/<int:kelas_id>')
def get_mapel_by_guru_and_kelas(guru_id, kelas_id):
    mapel_list = Mapel.query.filter_by(id_guru=guru_id, id_kelas=kelas_id).all()
    return jsonify([{'id': mapel.id_mapel, 'name': mapel.nama_mapel} for mapel in mapel_list])

@app.route('/get-semester-by-kelas/<int:kelas_id>')
def get_semester_by_kelas(kelas_id):
    semester_list = db.session.query(Semester).join(NilaiAkhir).filter(NilaiAkhir.id_kelas == kelas_id).distinct().all()
    return jsonify([{'id': semester.id, 'name': f"{semester.tahun_ajaran} - {semester.semester}"} for semester in semester_list])


@app.route('/download_nilai_excel', methods=['POST'])
@login_required
def download_nilai_excel():
    kelas_id = request.form['kelas']
    guru_id = request.form['guru']
    mapel_id = request.form['mapel']
    semester_id = request.form['semester']

    # Ambil data kelas, guru, mapel, dan semester
    kelas = Kelas.query.get(kelas_id)
    guru = Guru.query.get(guru_id)
    mapel = Mapel.query.get(mapel_id)
    semester = Semester.query.get(semester_id)

    # Ambil nilai akhir berdasarkan filter yang dipilih
    nilai_akhir_list = db.session.query(NilaiAkhir).join(Siswa).filter(
        NilaiAkhir.id_kelas == kelas_id,
        NilaiAkhir.id_guru == guru_id,
        NilaiAkhir.id_mapel == mapel_id,
        NilaiAkhir.id_semester == semester_id
    ).all()

    # Dictionary untuk mengelompokkan data berdasarkan guru dan mata pelajaran
    grouped_data = {}

    for nilai_akhir in nilai_akhir_list:
        siswa = nilai_akhir.siswa  # Mengambil objek siswa terkait
        if not siswa:
            continue

        guru_mapel_key = (guru.nama_guru, mapel.nama_mapel)

        if guru_mapel_key not in grouped_data:
            grouped_data[guru_mapel_key] = []

        grouped_data[guru_mapel_key].append((siswa.nama, nilai_akhir.nilai, nilai_akhir.capaian_kompetensi))

    # Buat file Excel
    wb = Workbook()
    ws = wb.active
    ws.append(['Nama Siswa', 'Tahun Ajaran', 'Semester', 'Nama Mapel', 'Nama Guru', 'Nilai', 'Capaian Kompetensi'])

    for (nama_guru, nama_mapel), entries in grouped_data.items():
        for entry in entries:
            nama_siswa, nilai, capaian_kompetensi = entry
            ws.append([nama_siswa, semester.tahun_ajaran, semester.semester, nama_mapel, nama_guru, nilai, capaian_kompetensi])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"{kelas.nama_kelas}_{guru.nama_guru}_{mapel.nama_mapel}_{semester.tahun_ajaran}_{semester.semester}.xlsx")

@app.route('/bulk_delete_and_download_nilai', methods=['GET', 'POST'])
@login_required
def bulk_delete_and_download_nilai():
    kelas = Kelas.query.order_by(Kelas.tingkat).all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran).all()
    
    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        mapel_id = request.form.get('mapel')
        semester_id = request.form.get('semester')
        
        if kelas_id and guru_id and mapel_id and semester_id:
            NilaiAkhir.query.filter_by(id_kelas=kelas_id, id_guru=guru_id, id_mapel=mapel_id, id_semester=semester_id).delete()
            db.session.commit()
            flash('Data nilai akhir berhasil dihapus.', 'success')
        else:
            flash('Lengkapi semua pilihan.', 'danger')

        # Redirect to the same page to clear form data
        return redirect(url_for('bulk_delete_and_download_nilai'))

    return render_template('bulk_delete_and_download_nilai.html', kelas=kelas, semesters=semesters)
    
    
@app.route('/bulk_delete_and_download_nilai_v2', methods=['GET', 'POST'])
@login_required
def bulk_delete_and_download_nilai_v2():
    kelas = Kelas.query.order_by(Kelas.tingkat).all()
    semesters = Semester.query.order_by(Semester.tahun_ajaran).all()
    
    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        guru_id = request.form.get('guru')
        mapel_id = request.form.get('mapel')
        semester_id = request.form.get('semester')
        
        if kelas_id and guru_id and mapel_id and semester_id:
            NilaiAkhir.query.filter_by(id_kelas=kelas_id, id_guru=guru_id, id_mapel=mapel_id, id_semester=semester_id).delete()
            db.session.commit()
            flash('Data nilai akhir berhasil dihapus.', 'success')
        else:
            flash('Lengkapi semua pilihan.', 'danger')

        # Redirect to the same page to clear form data
        return redirect(url_for('bulk_delete_and_download_nilai_v2'))

    return render_template('bulk_delete_and_download_nilai_v2.html', kelas=kelas, semesters=semesters)


@app.route('/download_nilai_excel_v2', methods=['POST'])
def download_nilai_excel_v2():
    kelas_id = request.form['kelas']
    guru_id = request.form['guru']
    mapel_id = request.form['mapel']
    semester_id = request.form['semester']

    kelas = Kelas.query.get(kelas_id)
    siswa_list = Siswa.query.filter_by(id_kelas=kelas_id).all()
    mapel = Mapel.query.get(mapel_id)
    guru = Guru.query.get(guru_id)

    wb = Workbook()
    ws = wb.active
    ws.append(['Nama Siswa', 'Nilai', 'Capaian Kompetensi'])

    for siswa in siswa_list:
        nilai_akhir = NilaiAkhir.query.filter_by(
            id_siswa=siswa.id_siswa,
            id_mapel=mapel_id,
            id_guru=guru_id,
            id_semester=semester_id
        ).first()
        
        if nilai_akhir:
            nilai = nilai_akhir.nilai
            capaian_kompetensi = nilai_akhir.capaian_kompetensi
        else:
            nilai = 'Belum dinilai'
            capaian_kompetensi = ''

        ws.append([siswa.nama, nilai, capaian_kompetensi])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name=f"{kelas.nama_kelas}_{mapel.nama_mapel}_{semester_id}.xlsx")
    
    
    
@app.route('/backup_database', methods=['GET'])
@login_required
def backup_database():
    try:
        backup_dir = 'backups'
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)

        backup_filename = f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
        backup_filepath = os.path.join(backup_dir, backup_filename)

        db_filepath = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        shutil.copy(db_filepath, backup_filepath)

        return send_file(backup_filepath, as_attachment=True, download_name=backup_filename)
    except Exception as e:
        flash(f"Backup failed: {str(e)}", "danger")
        return redirect(url_for('dashboard'))
    
@app.route('/restore_database', methods=['GET', 'POST'])
@login_required
def restore_database():
    if request.method == 'POST':
        file = request.files['backup_file']
        if file and file.filename.endswith('.db'):
            backup_dir = 'backups'
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)

            backup_filepath = os.path.join(backup_dir, file.filename)
            file.save(backup_filepath)

            db_filepath = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
            shutil.copy(backup_filepath, db_filepath)

            flash("Database restored successfully.", "success")
            return redirect(url_for('dashboard'))
        else:
            flash("Invalid file format. Please upload a .db file.", "danger")
            return redirect(url_for('restore_database'))

    return render_template('restore_database.html')


@app.route('/download_charts_pdf', methods=['GET'])
@login_required
def download_charts_pdf():
    class_id = request.args.get('classId')
    semester_id = request.args.get('semesterId')
    student_id = request.args.get('studentId')

    # Assuming you have a template 'charts_pdf.html' to render the charts
    rendered = render_template('charts_pdf.html', class_id=class_id, semester_id=semester_id, student_id=student_id)
    
    pdf = pdfkit.from_string(rendered, False)

    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=charts.pdf'
    return response

@app.route('/delete_rekap_absensi', methods=['GET', 'POST'])
@login_required
def delete_rekap_absensi():
    kelas_list = Kelas.query.order_by(Kelas.tingkat).all()
    semester_list = Semester.query.order_by(Semester.tahun_ajaran, Semester.semester).all()
    rekap_list = []

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        semester_id = request.form.get('semester')

        if 'confirm_delete' in request.form:
            # Perform the delete action
            rekap_to_delete = db.session.query(RekapAbsensi).join(Siswa).filter(
                Siswa.id_kelas == kelas_id,
                RekapAbsensi.id_semester == semester_id
            ).all()

            for rekap in rekap_to_delete:
                db.session.delete(rekap)
            db.session.commit()
            
            flash('Rekap Absensi berhasil dihapus.', 'success')
            return redirect(url_for('delete_rekap_absensi'))

        if kelas_id and semester_id:
            # Fetch the rekap_absensi records for display
            rekap_list = db.session.query(RekapAbsensi, Siswa)\
                .join(Siswa, RekapAbsensi.id_siswa == Siswa.id_siswa)\
                .filter(Siswa.id_kelas == kelas_id, RekapAbsensi.id_semester == semester_id)\
                .all()
        else:
            flash('Pilih Kelas dan Semester.', 'danger')

    return render_template('delete_rekap_absensi.html', kelas_list=kelas_list, semester_list=semester_list, rekap_list=rekap_list)
 