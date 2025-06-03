from flask import current_app, redirect, render_template, request, Blueprint, send_file, send_from_directory,jsonify, url_for
from zipfile import ZipFile
from app import db
from flask_login import login_required, current_user
from functools import wraps
from app.models import Kelas, Siswa, Semester, NilaiAkhir, Pengajaran, Mapel, Ekstrakurikuler, RekapAbsensi, KegiatanIndustri
from docx import Document
from docx.shared import Pt, Cm, Inches  # Import Pt for point size
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # Import enum for paragraph alignment
from docx.enum.table import WD_ALIGN_VERTICAL
import io
import tempfile
import os



generate_bp = Blueprint('generate', __name__, url_prefix='/generate-raporsiswa')

@generate_bp.route('/generate-raporsiswa', methods=['GET', 'POST'])
@login_required
def generate_raporsiswa():
    kelas_list = Kelas.query.order_by(Kelas.nama_kelas.asc()).all()
    siswa_list = Siswa.query.order_by(Siswa.nama.asc()).all()
   # Ambil daftar semester dengan urutan ascending berdasarkan semester
    semester_list = Semester.query.options(db.joinedload(Semester.tahun_ajaran)).order_by(Semester.semester.asc()).all()

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        siswa_id = request.form.get('siswa')
        semester_id = request.form.get('semester')

        if kelas_id and siswa_id and semester_id:
            siswa = Siswa.query.filter_by(id_siswa=siswa_id).first()

            # Query NilaiAkhir dengan join ke Pengajaran dan Mapel
            nilai_akhir = db.session.query(NilaiAkhir, Mapel).join(Pengajaran, NilaiAkhir.id_pengajaran == Pengajaran.id_pengajaran)\
                .join(Mapel, Pengajaran.id_mapel == Mapel.id_mapel)\
                .filter(NilaiAkhir.id_siswa == siswa_id, Pengajaran.id_semester == semester_id).all()

            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            semester = Semester.query.filter_by(id=semester_id).first()

            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katholik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Matematika": 5,
                "Fisika": 6,
                "Kimia": 7,
                "Biologi": 8,
                "Sosiologi": 9,
                "Ekonomi": 10,
                "Sejarah": 11,
                "Geografi": 12,
                "Bahasa Inggris": 13,
                "Pendidikan Jasmani Olahraga dan Kesehatan": 14,
                "Informatika": 15,
                "Seni Tari": 16,
                "Bahasa Mandarin": 17,
                "Teknologi Informasi dan Komunikasi (TIK)": 18,
            }
 
            return render_template('rapor_siswa.html', siswa=siswa, nilai_akhir=nilai_akhir, ekstrakurikuler=ekstrakurikuler, 
                                   rekap_absensi=rekap_absensi, kegiatan_industri=kegiatan_industri, 
                                   mapel_ke_baris=mapel_ke_baris, kelas_list=kelas_list, siswa_list=siswa_list, 
                                   semester_list=semester_list, semester=semester)

    return render_template('rapor_siswa.html', kelas_list=kelas_list, siswa_list=siswa_list, semester_list=semester_list)

# Add custom decorator to check authentication
def check_auth(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            return jsonify({'error': 'Unauthorized'}), 401
        return f(*args, **kwargs)
    return decorated_function

@generate_bp.route('/generate-rapor-kelas/<int:id_kelas>/<int:id_semester>')
@login_required
@check_auth  # Add extra authentication check
def generate_rapor_kelas(id_kelas, id_semester):
    try:
        # Add session check
        if not current_user.is_authenticated:
            return redirect(url_for('auth.login'))
            
        kelas = Kelas.query.get(id_kelas)
        if kelas is None:
            return "Kelas tidak ditemukan", 404

        nama_kelas = kelas.nama_kelas
        if nama_kelas == "X":
            response = generate_rapor_kelas_x(id_kelas, id_semester)
            # Add security headers
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            return response
        else:
            return "Kelas tidak dikenal", 400
            
    except Exception as e:
        current_app.logger.error(f"Error generating report: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

# raport by kelas


@generate_bp.route('/generate-rapor-x/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_x(id_kelas, id_semester):
    # Ambil semua siswa dalam kelas
    kelas = Kelas.query.get(id_kelas)
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()

    # ... sisa kode

            
            mapel_ke_baris = {
                    "Pendidikan Agama Islam dan Budi Pekerti": 2,
                    "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                    "Pendidikan Agama Katholik dan Budi Pekerti": 2,
                    "Pendidikan Pancasila": 3,
                    "Bahasa Indonesia": 4,
                    "Matematika": 5,
                    "Fisika": 6,
                    "Kimia": 7,
                    "Biologi": 8,
                    "Sosiologi": 9,
                    "Ekonomi": 10,
                    "Sejarah": 11,
                    "Geografi": 12,
                    "Bahasa Inggris": 13,
                    "Pendidikan Jasmani Olahraga dan Kesehatan": 14,
                    "Informatika": 15,
                    "Seni Tari": 16,
                    "Bahasa Mandarin": 18,
                    "Teknologi Informasi dan Komunikasi (TIK)": 19,
            }

            # Buka template
            doc = Document('template_rapor_x.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8

            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)
                counter += 1  # Move increment inside loop

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]
            section.footer_distance = Cm(1.0)  # Set footer 1cm from bottom

            # Create or get footer
            footer = section.footer

            # Clear existing footer content
            for paragraph in footer.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
                paragraph._p = paragraph._element = None

            # Add new footer with proper formatting
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.size = Pt(10)
            footer_para.style.font.name = 'Times New Roman'

            # Add footer text
            footer_text = f"Nama: {siswa.nama} | NISN: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            run = footer_para.add_run(footer_text)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

            # Add bottom border to footer
            footer_para.paragraph_format.space_after = Pt(0)
            footer_para.paragraph_format.space_before = Pt(0)

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

# Fungsi untuk mengatur gaya sel tabel


def atur_gaya_sel(sel, nama_font='Times New Roman', ukuran_font=12, tengah_align=False, justify_align=False):
    """
    Mengatur gaya untuk sel dalam dokumen.

    Parameters:
    - sel: Sel yang akan diatur gayanya.
    - nama_font: Nama font yang akan digunakan (default: 'Times New Roman').
    - ukuran_font: Ukuran font yang akan digunakan (default: 12).
    - tengah_align: Jika True, maka paragraf akan diatur menjadi rata tengah (default: False).
    - justify_align: Jika True, maka paragraf akan diatur menjadi rata kanan dan kiri (default: False).
    """
    sel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraf in sel.paragraphs:
        # Atur alignment
        if tengah_align:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif justify_align:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Atur font
        for run in paragraf.runs:
            run.font.name = nama_font
            run.font.size = Pt(ukuran_font)

# Pastikan juga untuk mengimpor WD_ALIGN_PARAGRAPH dari docx.enum.text

def atur_tinggi_baris(tabel, tinggi_cm):
    for baris in tabel.rows:
        for sel in baris.cells:
            sel.height = Cm(tinggi_cm)

from docx.shared import Cm

# Fungsi untuk mengatur tinggi baris dalam tabel
def atur_tinggi_baris(tabel, baris, tinggi_cm):
    for sel in tabel.rows[baris].cells:
        sel.height = Cm(tinggi_cm)


@generate_bp.route('/get_siswa_by_kelas/<int:kelas_id>', methods=['GET'])
def get_siswa_by_kelas(kelas_id):
    siswa_list = Siswa.query.filter_by(id_kelas=kelas_id).order_by(Siswa.nama.asc()).all()
    data = [{'id_siswa': s.id_siswa, 'nama': s.nama} for s in siswa_list]
    return jsonify(data)


