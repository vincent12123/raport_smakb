from flask import render_template, request, Blueprint, send_file, send_from_directory,jsonify
from zipfile import ZipFile
from app import db
from flask_login import login_required
from app.models import Kelas, Siswa, Semester, NilaiAkhir, Pengajaran, Mapel, Ekstrakurikuler, RekapAbsensi, KegiatanIndustri
from docx import Document
from docx.shared import Pt, Cm  # Import Pt for point size
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # Import enum for paragraph alignment
from docx.enum.table import WD_ALIGN_VERTICAL
import io
import tempfile
import os



generate_bp = Blueprint('generate', __name__, url_prefix='/generate-raporsiswa')

@generate_bp.route('/generate-raporsiswa', methods=['GET', 'POST'])
@login_required
def generate_raporsiswa():
    kelas_list = Kelas.query.order_by(Kelas.nama_kelas.asc()).all()
    siswa_list = Siswa.query.order_by(Siswa.nama.asc()).all()
   # Ambil daftar semester dengan urutan ascending berdasarkan semester
    semester_list = Semester.query.options(db.joinedload(Semester.tahun_ajaran)).order_by(Semester.semester.asc()).all()

    if request.method == 'POST':
        kelas_id = request.form.get('kelas')
        siswa_id = request.form.get('siswa')
        semester_id = request.form.get('semester')

        if kelas_id and siswa_id and semester_id:
            siswa = Siswa.query.filter_by(id_siswa=siswa_id).first()

            # Query NilaiAkhir dengan join ke Pengajaran dan Mapel
            nilai_akhir = db.session.query(NilaiAkhir, Mapel).join(Pengajaran, NilaiAkhir.id_pengajaran == Pengajaran.id_pengajaran)\
                .join(Mapel, Pengajaran.id_mapel == Mapel.id_mapel)\
                .filter(NilaiAkhir.id_siswa == siswa_id, Pengajaran.id_semester == semester_id).all()

            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa_id, id_semester=semester_id).all()
            semester = Semester.query.filter_by(id=semester_id).first()

            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,
                "Seni Musik": 7,
                "Matematika": 8,
                "Bahasa Inggris": 9,
                "Rekayasa Perangkat Lunak": 10,
                "Teknik Sepeda Motor": 10,
                "Perhotelan": 10,
                "Projek Kreatif dan Kewirausahaan": 11,
                "Informatika": 11,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 12,
                "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 13,
                "Digital Marketing": 12,
                "Dasar-Dasar Otomotif": 13,
                "Dasar-Dasar Perhotelan": 13,
            }

            return render_template('rapor_siswa.html', siswa=siswa, nilai_akhir=nilai_akhir, ekstrakurikuler=ekstrakurikuler, 
                                   rekap_absensi=rekap_absensi, kegiatan_industri=kegiatan_industri, 
                                   mapel_ke_baris=mapel_ke_baris, kelas_list=kelas_list, siswa_list=siswa_list, 
                                   semester_list=semester_list, semester=semester)

    return render_template('rapor_siswa.html', kelas_list=kelas_list, siswa_list=siswa_list, semester_list=semester_list)

@generate_bp.route('/generate-rapor-kelas/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas(id_kelas, id_semester):
    # Retrieve class data from the database
    kelas = Kelas.query.get(id_kelas)
    
    # Check if the class was found
    if kelas is None:
        return "Kelas tidak ditemukan", 404
    
    nama_kelas = kelas.nama_kelas

    # Determine which report generation function to call based on the class and semester
    if nama_kelas == "XI Rekayasa Perangkat Lunak":
        return generate_rapor_kelas_rplxi(id_kelas, id_semester)
    elif nama_kelas == "XI Teknik Sepeda Motor":
        return generate_rapor_kelas_tsmxi(id_kelas, id_semester)
    elif nama_kelas == "XI Perhotelan":
        return generate_rapor_kelas_htlxi(id_kelas, id_semester)
    elif nama_kelas == "X Rekayasa Perangkat Lunak":
        return generate_rapor_kelas_rplx(id_kelas, id_semester)
    elif nama_kelas == "X Teknik Sepeda Motor":
        return generate_rapor_kelas_tsmx(id_kelas, id_semester)
    elif nama_kelas == "X Perhotelan":
        return generate_rapor_kelas_htlx(id_kelas, id_semester)
    else:
        return "Kelas tidak dikenal", 400

# raport by kelas

@generate_bp.route('/generate-rapor-rplxi/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_rplxi(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)

    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        for siswa in siswa_dalam_kelas:
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()

            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()

            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,
                "Matematika": 9,
                "Bahasa Inggris": 10,
                "Rekayasa Perangkat Lunak": 11,
                "Projek Kreatif dan Kewirausahaan": 12,
                "Digital Marketing": 13,
            }

            doc = Document('template_rapor_newxi_rpl.docx')   

            # Isi data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)     

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel

                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)

                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi        

            # Isi kegiatan ekstrakurikuler
            counter = 1
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                baris.cells[0].text = str(counter)
                baris.cells[1].text = ekstra.kegiatan
                baris.cells[2].text = ekstra.keterangan
                counter += 1    

            # Isi kegiatan industri
            counter = 1  
            for kegiatan in kegiatan_industri:
                baris = doc.tables[2].add_row()
                baris.cells[0].text = str(counter)
                baris.cells[1].text = kegiatan.mitra_induka
                baris.cells[2].text = kegiatan.lokasi
                baris.cells[3].text = kegiatan.keterangan
                counter += 1

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                doc.tables[3].rows[0].cells[1].text = str(rekap_absensi.total_sakit) + ' Hari'
                doc.tables[3].rows[1].cells[1].text = str(rekap_absensi.total_izin) + ' Hari'
                doc.tables[3].rows[2].cells[1].text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
            
            # Tambahkan footer
            section = doc.sections[0]
            footer = section.footer
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer.paragraphs[0].text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

@generate_bp.route('/generate-rapor-tsmxi/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_tsmxi(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)

    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()
            
            
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
    # ... sisa kode
           # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,

                "Matematika": 9,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 10,
                "Teknik Sepeda Motor": 11,
                "Projek Kreatif dan Kewirausahaan": 12,
                "Digital Marketing": 13,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_newxi_tsm.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan, justify_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter, tinggi_cm)

                counter += 1    

            # Isi kegiatan industri
            counter = 1  
            for kegiatan in kegiatan_industri:
                baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
                sel_nomor = baris.cells[0]  # Sel nomor
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_mitra_industri = baris.cells[1]  # Sel mitra industri
                sel_mitra_industri.text = kegiatan.mitra_induka
                atur_gaya_sel(sel_mitra_industri)

                sel_lokasi = baris.cells[2]  # Sel lokasi
                sel_lokasi.text = kegiatan.lokasi
                atur_gaya_sel(sel_lokasi, tengah_align=True)

                sel_keterangan = baris.cells[3]  # Sel keterangan
                sel_keterangan.text = kegiatan.keterangan
                atur_gaya_sel(sel_keterangan, tengah_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[2], counter, tinggi_cm)

                counter += 1

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[3].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[3].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[3].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            
            # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

    #return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{id_kelas}.zip')

@generate_bp.route('/generate-rapor-htlxi/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_htlxi(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporanf
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()
            kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Mandarin": 7,

                "Matematika": 9,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 10,
                "Perhotelan": 11,
                "Projek Kreatif dan Kewirausahaan": 12,
                "Digital Marketing": 13,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_newxi_htl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)
            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan, justify_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter, tinggi_cm)

                counter += 1    

            # Isi kegiatan industri
            counter = 1  
            for kegiatan in kegiatan_industri:
                baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
                sel_nomor = baris.cells[0]  # Sel nomor
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_mitra_industri = baris.cells[1]  # Sel mitra industri
                sel_mitra_industri.text = kegiatan.mitra_induka
                atur_gaya_sel(sel_mitra_industri)

                sel_lokasi = baris.cells[2]  # Sel lokasi
                sel_lokasi.text = kegiatan.lokasi
                atur_gaya_sel(sel_lokasi, tengah_align=True)

                sel_keterangan = baris.cells[3]  # Sel keterangan
                sel_keterangan.text = kegiatan.keterangan
                atur_gaya_sel(sel_keterangan, tengah_align=True)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[2], counter, tinggi_cm)

                counter += 1

            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[3].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[3].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[3].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            
            # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

@generate_bp.route('/generate-rapor-rplx/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_rplx(id_kelas, id_semester):
    # Ambil semua siswa dalam kelas
    kelas = Kelas.query.get(id_kelas)
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()

    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Seni Musik": 7,
                "Mandarin": 8,

                "Matematika": 10,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 11,
                "Informatika": 12,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
                "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 14,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_x_rpl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)

            counter += 1    
            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')


@generate_bp.route('/generate-rapor-tsmx/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_tsmx(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).first()

    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Seni Musik": 7,
                "Mandarin": 8,

                "Matematika": 10,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 11,
                "Informatika": 12,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
                "Dasar-Dasar Otomotif": 14,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_x_tsm.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)

            counter += 1    
            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')


@generate_bp.route('/generate-rapor-htlx/<int:id_kelas>/<int:id_semester>')
@login_required
def generate_rapor_kelas_htlx(id_kelas, id_semester):
    kelas = Kelas.query.get(id_kelas)
    # Ambil semua siswa dalam kelas
    siswa_dalam_kelas = Siswa.query.filter_by(id_kelas=id_kelas).all()
    semester = Semester.query.get(id_semester)
    # Buat file zip sementara
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'a') as zip_file:
        # Loop melalui setiap siswa dan buat laporan
        for siswa in siswa_dalam_kelas:
    # Menggunakan siswa.id sebagai pengganti id_siswa
            nilai_akhir = db.session.query(NilaiAkhir).join(Pengajaran).filter(
                NilaiAkhir.id_siswa == siswa.id_siswa,
                Pengajaran.id_semester == id_semester
            ).all()
            ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=siswa.id_siswa, id_semester=id_semester).all()
            rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=siswa.id_siswa,id_semester=id_semester).first()

    # ... sisa kode

            
            # Pemetaan nama mapel ke nomor baris di tabel
            mapel_ke_baris = {
                "Pendidikan Agama Islam dan Budi Pekerti": 2,
                "Pendidikan Agama Kristen dan Budi Pekerti": 2,
                "Pendidikan Agama Katolik dan Budi Pekerti": 2,
                "Pendidikan Pancasila": 3,
                "Bahasa Indonesia": 4,
                "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
                "Sejarah": 6,
                "Seni Musik": 7,
                "Mandarin": 8,

                "Matematika": 10,  # Misalnya Matematika ada di baris pertama
                "Bahasa Inggris": 11,
                "Informatika": 12,
                "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
                "Dasar-Dasar Perhotelan": 14,  # Bahasa Inggris di baris kedua
                
                
                # Tambahkan pemetaan untuk mata pelajaran lainnya
            }

            # Buka template
            doc = Document('template_rapor_x_htl.docx')   
            
            # Isi bagian data siswa
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '{Nama}' in run.text:
                        run.text = run.text.replace('{Nama}', siswa.nama)
                    if '{Kelas}' in run.text:
                        run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                    if '{NISN/NIS}' in run.text:
                        run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
                    if '{Semester}' in run.text:
                        run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                    if '{TahunAjaran}' in run.text:
                        run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran.tahun)

            # Isi nilai akhir
            for nilai in nilai_akhir:
                baris = mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
                if baris is not None:
                    # Sel nama mata pelajaran
                    sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
                    sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
                    atur_gaya_sel(sel_nama_mapel)

                    # Sel nilai
                    sel_nilai = doc.tables[0].rows[baris].cells[2]
                    sel_nilai.text = str(nilai.nilai)
                    atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

                    # Sel capaian kompetensi
                    sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
                    sel_capaian_kompetensi.text = nilai.capaian_kompetensi
                    atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)        
            
            # Isi kegiatan ekstrakurikuler
            counter = 1
            tinggi_cm = 0.8  # Tinggi baris dalam cm
            for ekstra in ekstrakurikuler:
                baris = doc.tables[1].add_row()
                sel_nomor = baris.cells[0]
                sel_nomor.text = str(counter)
                atur_gaya_sel(sel_nomor, tengah_align=True)

                sel_kegiatan = baris.cells[1]
                sel_kegiatan.text = ekstra.kegiatan
                atur_gaya_sel(sel_kegiatan)

                sel_keterangan = baris.cells[2]
                sel_keterangan.text = ekstra.keterangan
                atur_gaya_sel(sel_keterangan)

                # Mengatur tinggi baris
                atur_tinggi_baris(doc.tables[1], counter - 1, tinggi_cm)

            counter += 1    
            # Isi rekapitulasi kehadiran
            if rekap_absensi:
                # Sel total sakit
                sel_total_sakit = doc.tables[2].rows[0].cells[1]
                sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
                atur_gaya_sel(sel_total_sakit)

                # Sel total izin
                sel_total_izin = doc.tables[2].rows[1].cells[1]
                sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
                atur_gaya_sel(sel_total_izin)

                # Sel total tanpa keterangan
                sel_total_tanpa_keterangan = doc.tables[2].rows[2].cells[1]
                sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
                atur_gaya_sel(sel_total_tanpa_keterangan)
            # Isi rekapitulasi kehadiran
            
          
        # Ambil bagian atas halaman
            section = doc.sections[0]

            # Memanggil footer
            footer = section.footer

            # Tambahkan teks footer dengan informasi siswa
            footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text

            # Simpan dokumen ke dalam file sementara
            with tempfile.NamedTemporaryFile(delete=False) as temp:
                doc.save(temp.name)
                temp.close()

                # Tambahkan file sementara ke file zip
                zip_file.write(temp.name, f'rapor_{siswa.nama}.docx')

                # Hapus file sementara
                os.unlink(temp.name)

    # Siapkan file zip untuk diunduh
    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'rapor_kelas_{kelas.nama_kelas}_semester_{id_semester}.zip')

# end ..

# Fungsi untuk mengatur gaya sel tabel


def atur_gaya_sel(sel, nama_font='Times New Roman', ukuran_font=12, tengah_align=False, justify_align=False):
    """
    Mengatur gaya untuk sel dalam dokumen.

    Parameters:
    - sel: Sel yang akan diatur gayanya.
    - nama_font: Nama font yang akan digunakan (default: 'Times New Roman').
    - ukuran_font: Ukuran font yang akan digunakan (default: 12).
    - tengah_align: Jika True, maka paragraf akan diatur menjadi rata tengah (default: False).
    - justify_align: Jika True, maka paragraf akan diatur menjadi rata kanan dan kiri (default: False).
    """
    sel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraf in sel.paragraphs:
        # Atur alignment
        if tengah_align:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif justify_align:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Atur font
        for run in paragraf.runs:
            run.font.name = nama_font
            run.font.size = Pt(ukuran_font)

# Pastikan juga untuk mengimpor WD_ALIGN_PARAGRAPH dari docx.enum.text

def atur_tinggi_baris(tabel, tinggi_cm):
    for baris in tabel.rows:
        for sel in baris.cells:
            sel.height = Cm(tinggi_cm)

from docx.shared import Cm

# Fungsi untuk mengatur tinggi baris dalam tabel
def atur_tinggi_baris(tabel, baris, tinggi_cm):
    for sel in tabel.rows[baris].cells:
        sel.height = Cm(tinggi_cm)


@generate_bp.route('/get_siswa_by_kelas/<int:kelas_id>', methods=['GET'])
def get_siswa_by_kelas(kelas_id):
    siswa_list = Siswa.query.filter_by(id_kelas=kelas_id).order_by(Siswa.nama.asc()).all()
    data = [{'id_siswa': s.id_siswa, 'nama': s.nama} for s in siswa_list]
    return jsonify(data)


# Generate by name
@generate_bp.route('/generate-rapor-nama')
@login_required
def generate_rapor_nama():
    nama_siswa = request.args.get('nama_siswa')
    # Retrieve student data and grades from the database based on the student's name
   
    siswa = Siswa.query.filter_by(nama=nama_siswa).first()
    
    # Check if the student was found
    if siswa is None:
        return "Siswa tidak ditemukan", 404
    id_siswa = siswa.id_siswa
    id_kelas = siswa.id_kelas
    nama_kelas = siswa.kelas.nama_kelas

    # Determine which report generation function to call based on the student's class
    if nama_kelas == "XI Rekayasa Perangkat Lunak":
        #return generate_rapor_kelas_rplxi(id_kelas)
        return generate_rapor_rplxi(id_siswa)
    elif nama_kelas == "XI Teknik Sepeda Motor":
        #return generate_rapor_kelas_tsmxi(id_kelas)
        return generate_rapor_tsmxi(id_siswa)
    elif nama_kelas == "XI Perhotelan":
        #return generate_rapor_kelas_htlxi(id_kelas)
        return generate_rapor_htlxi(id_siswa)
    elif nama_kelas == "X Rekayasa Perangkat Lunak":
        #return generate_rapor_kelas_rplx(id_kelas)
        return generate_rapor_rplx(id_siswa)
    elif nama_kelas == "X Teknik Sepeda Motor":
        #return generate_rapor_kelas_tsmx(id_kelas)
        return generate_rapor_tsmx(id_siswa)
    elif nama_kelas == "X Perhotelan":
        #return generate_rapor_kelas_htlx(id_kelas)
        return generate_rapor_htlx(id_siswa)
    else:
        return "Kelas tidak dikenal", 400


def generate_rapor_rplxi(id_siswa):
    mapel_ke_baris_rplxi = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Mandarin": 7,
        "Matematika": 9,
        "Bahasa Inggris": 10,
        "Rekayasa Perangkat Lunak": 11,
        "Projek Kreatif dan Kewirausahaan": 12,
        "Digital Marketing": 13,
    }
    return generate_rapor(id_siswa, 'template_rapor_newxi_rpl.docx', mapel_ke_baris_rplxi, include_kegiatan_industri=True)

def generate_rapor_tsmxi(id_siswa):
    mapel_ke_baris_tsmxi = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Mandarin": 7,
        "Matematika": 9,
        "Bahasa Inggris": 10,
        "Teknik Sepeda Motor": 11,
        "Projek Kreatif dan Kewirausahaan": 12,
        "Digital Marketing": 13,
    }
    return generate_rapor(id_siswa, 'template_rapor_newxi_tsm.docx', mapel_ke_baris_tsmxi, include_kegiatan_industri=True)

def generate_rapor_htlxi(id_siswa):
    mapel_ke_baris_htlxi = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Mandarin": 7,
        "Matematika": 9,
        "Bahasa Inggris": 10,
        "Perhotelan": 11,
        "Projek Kreatif dan Kewirausahaan": 12,
        "Digital Marketing": 13,
    }
    return generate_rapor(id_siswa, 'template_rapor_newxi_htl.docx', mapel_ke_baris_htlxi, include_kegiatan_industri=True)

def generate_rapor_rplx(id_siswa):
    mapel_ke_baris_rplx = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Seni Musik": 7,
        "Mandarin": 8,
        "Matematika": 10,  
        "Bahasa Inggris": 11,
        "Informatika": 12,
        "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
        "Dasar-Dasar Rekayasa Perangkat Lunak & Gim": 14,
    }
    return generate_rapor(id_siswa, 'template_rapor_x_rpl.docx', mapel_ke_baris_rplx)

def generate_rapor_tsmx(id_siswa):
    mapel_ke_baris_tsmx = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Seni Musik": 7,
        "Mandarin": 8,
        "Matematika": 10,  
        "Bahasa Inggris": 11,
        "Informatika": 12,
        "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
        "Dasar-Dasar Otomotif": 14,
    }
    return generate_rapor(id_siswa, 'template_rapor_x_tsm.docx', mapel_ke_baris_tsmx)

def generate_rapor_htlx(id_siswa):
    mapel_ke_baris_htlx = {
        "Pendidikan Agama Islam dan Budi Pekerti": 2,
        "Pendidikan Agama Kristen dan Budi Pekerti": 2,
        "Pendidikan Agama Katolik dan Budi Pekerti": 2,
        "Pendidikan Pancasila": 3,
        "Bahasa Indonesia": 4,
        "Pendidikan Jasmani, Olahraga, dan Kesehatan": 5,
        "Sejarah": 6,
        "Seni Musik": 7,
        "Mandarin": 8,
        "Matematika": 10,  
        "Bahasa Inggris": 11,
        "Informatika": 12,
        "Projek Ilmu Pengetahuan Alam dan Sosial": 13,
        "Dasar-Dasar Perhotelan": 14,
    }
    return generate_rapor(id_siswa, 'template_rapor_x_htl.docx', mapel_ke_baris_htlx)

# end form rapor persiswa
def generate_rapor(id_siswa, template_path, mapel_ke_baris, include_kegiatan_industri=False):
    # Ambil data siswa dan nilai dari database
    siswa = Siswa.query.get(id_siswa)
    nilai_akhir = NilaiAkhir.query.filter_by(id_siswa=id_siswa).all()
    ekstrakurikuler = Ekstrakurikuler.query.filter_by(id_siswa=id_siswa).all()
    rekap_absensi = RekapAbsensi.query.filter_by(id_siswa=id_siswa).first()
   
    # Periksa jika kegiatan_industri perlu dimasukkan berdasarkan template
    if "rplxi" in template_path or "tsmxi" in template_path or "htlxi" in template_path or include_kegiatan_industri:
        kegiatan_industri = KegiatanIndustri.query.filter_by(id_siswa=id_siswa).all()
    else:
        kegiatan_industri = [] 

    # Ambil informasi semester dari salah satu nilai akhir
    if nilai_akhir:
        semester = Semester.query.get(nilai_akhir[0].id_semester)
    else:
        semester = None

    # Buka template
    doc = Document(template_path)
    
    # Isi bagian data siswa
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '{Nama}' in run.text and siswa.nama is not None:
                run.text = run.text.replace('{Nama}', siswa.nama)
            if '{Kelas}' in run.text:
                if siswa.kelas.nama_kelas is not None:
                    run.text = run.text.replace('{Kelas}', siswa.kelas.nama_kelas)
                else:
                    run.text = run.text.replace('{Kelas}', 'Default Class')
            if '{NISN/NIS}' in run.text and siswa.nisn is not None:
                run.text = run.text.replace('{NISN/NIS}', siswa.nisn)
            if semester:
                if '{Semester}' in run.text:
                    run.text = run.text.replace('{Semester}', f"{semester.semester} / {semester.fase}")
                if '{TahunAjaran}' in run.text:
                    run.text = run.text.replace('{TahunAjaran}', semester.tahun_ajaran)

    # Isi nilai akhir
    for nilai in nilai_akhir:
        mapel_ke_baris.get(nilai.pengajaran.mapel.nama_mapel)
        if baris is not None:
            # Sel nama mata pelajaran
            sel_nama_mapel = doc.tables[0].rows[baris].cells[1]
            sel_nama_mapel.text = nilai.pengajaran.mapel.nama_mapel
            atur_gaya_sel(sel_nama_mapel)

            # Sel nilai
            sel_nilai = doc.tables[0].rows[baris].cells[2]
            sel_nilai.text = str(nilai.nilai)
            atur_gaya_sel(sel_nilai, tengah_align=True)  # Tengah-align teks

            # Sel capaian kompetensi
            sel_capaian_kompetensi = doc.tables[0].rows[baris].cells[3]
            sel_capaian_kompetensi.text = nilai.capaian_kompetensi
            atur_gaya_sel(sel_capaian_kompetensi, justify_align=True)

    # Isi kegiatan ekstrakurikuler
    counter = 1
    for ekstra in ekstrakurikuler:
        baris = doc.tables[1].add_row()
        sel_nomor = baris.cells[0]
        sel_nomor.text = str(counter)
        atur_gaya_sel(sel_nomor, tengah_align=True)

        sel_kegiatan = baris.cells[1]
        sel_kegiatan.text = ekstra.kegiatan
        atur_gaya_sel(sel_kegiatan)

        sel_keterangan = baris.cells[2]
        sel_keterangan.text = ekstra.keterangan
        atur_gaya_sel(sel_keterangan)

        counter += 1    

    # Isi kegiatan industri jika ada
    counter = 1
    if kegiatan_industri:
        for kegiatan in kegiatan_industri:
            baris = doc.tables[2].add_row()  # Tambahkan baris baru pada tabel kegiatan industri
            sel_nomor = baris.cells[0]  # Sel nomor
            sel_nomor.text = str(counter)
            atur_gaya_sel(sel_nomor, tengah_align=True)

            sel_mitra_industri = baris.cells[1]  # Sel mitra industri
            sel_mitra_industri.text = kegiatan.mitra_induka
            atur_gaya_sel(sel_mitra_industri)

            sel_lokasi = baris.cells[2]  # Sel lokasi
            sel_lokasi.text = kegiatan.lokasi
            atur_gaya_sel(sel_lokasi, tengah_align=True)

            sel_keterangan = baris.cells[3]  # Sel keterangan
            sel_keterangan.text = kegiatan.keterangan
            atur_gaya_sel(sel_keterangan, tengah_align=True)

            counter += 1    

    # Isi rekapitulasi kehadiran
    # Periksa jumlah tabel dalam dokumen
    table_index = 3 if kegiatan_industri else 2
    if rekap_absensi:
        # Sel total sakit
        sel_total_sakit = doc.tables[table_index].rows[0].cells[1]
        sel_total_sakit.text = str(rekap_absensi.total_sakit) + ' Hari'
        atur_gaya_sel(sel_total_sakit)

        # Sel total izin
        sel_total_izin = doc.tables[table_index].rows[1].cells[1]
        sel_total_izin.text = str(rekap_absensi.total_izin) + ' Hari'
        atur_gaya_sel(sel_total_izin)

        # Sel total tanpa keterangan
        sel_total_tanpa_keterangan = doc.tables[table_index].rows[2].cells[1]
        sel_total_tanpa_keterangan.text = str(rekap_absensi.total_tanpa_keterangan) + ' Hari'
        atur_gaya_sel(sel_total_tanpa_keterangan)

    # Ambil bagian atas halaman
    section = doc.sections[0]

    # Memanggil footer
    footer = section.footer

    # Tambahkan teks footer dengan informasi siswa
    footer_text = f"Nama: {siswa.nama} | NIS: {siswa.nisn} | Kelas: {siswa.kelas.nama_kelas}"
    footer_para = footer.paragraphs[0]
    footer_para.text = footer_text

    # Simpan dokumen ke dalam buffer dan kirim sebagai file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'rapor_{siswa.nama}.docx')

